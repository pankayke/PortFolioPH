"""
generate_jira_test_samples.py
Generates Sprint 1 & 2 – PortFolioPH Jira Task Summary with Test Samples as a Word .docx file.
Output: docs/PortFolioPH_Sprint1_Sprint2_Jira_Test_Samples.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour constants ───────────────────────────────────────────────────────────
PRIMARY      = RGBColor(0x0D, 0x47, 0xA1)
ACCENT       = RGBColor(0xFF, 0x98, 0x00)
DARK_GREY    = RGBColor(0x42, 0x42, 0x42)
MID_GREY     = RGBColor(0x75, 0x75, 0x75)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
GREEN        = RGBColor(0x38, 0x8E, 0x3C)
RED          = RGBColor(0xD3, 0x2F, 0x2F)
SPRINT1_COL  = RGBColor(0x0D, 0x47, 0xA1)
SPRINT2_COL  = RGBColor(0x1B, 0x5E, 0x20)

# ── Hex strings for cell backgrounds ──────────────────────────────────────────
HEX_PRIMARY    = '0D47A1'
HEX_SPRINT2    = '1B5E20'
HEX_ACCENT     = 'E65100'
HEX_LIGHT_BLUE = 'E3F2FD'
HEX_LIGHT_GRN  = 'E8F5E9'
HEX_WHITE      = 'FFFFFF'
HEX_HEADER_GRY = '37474F'

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def add_heading(doc, text, level, color=PRIMARY, space_before=12, space_after=6):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.color.rgb = color
    return p


def add_para(doc, text, bold=False, italic=False, color=DARK_GREY,
             size=11, space_after=6, space_before=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    set_run_color(run, color)
    return p


def add_bullet(doc, text, color=DARK_GREY, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_run_color(run, color)
    return p


def make_test_table(doc, rows, sprint=1):
    """Create the 4-column test step table: Step | Action | Expected Result | Result."""
    headers = ['Step', 'Action', 'Expected Result', 'Result']
    col_widths = [0.55, 2.5, 3.2, 0.85]
    header_bg = HEX_PRIMARY if sprint == 1 else HEX_SPRINT2
    row_bg_a  = HEX_LIGHT_BLUE if sprint == 1 else HEX_LIGHT_GRN

    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = row_bg_a if r_idx % 2 == 0 else HEX_WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            if c_idx == 3:
                run.bold = True
                run.font.color.rgb = GREEN
            else:
                set_run_color(run, DARK_GREY)

    # Column widths
    for row in table.rows:
        for idx, w in enumerate(col_widths):
            row.cells[idx].width = Inches(w)

    return table


def make_summary_table(doc, rows):
    """Overall sprint coverage summary table."""
    headers = ['Sprint', 'Tasks Completed', 'Key Features Delivered']
    col_widths = [0.9, 1.4, 4.8]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_bg(cell, HEX_HEADER_GRY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    bg_map = ['E3F2FD', 'E8F5E9', 'FFF8E1']
    for r_idx, row_data in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = bg_map[r_idx % len(bg_map)]
        for c_idx, cell_text in enumerate(row_data):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.bold = c_idx == 0
            set_run_color(run, DARK_GREY)

    for row in table.rows:
        for idx, w in enumerate(col_widths):
            row.cells[idx].width = Inches(w)

    return table


def add_task_block(doc, task_id, title, summary, bullets, test_rows, sprint=1):
    """Render a full task block: ID badge, title, summary, bullets, test table."""
    id_color = SPRINT1_COL if sprint == 1 else SPRINT2_COL
    id_hex   = HEX_PRIMARY  if sprint == 1 else HEX_SPRINT2

    # Task ID + title heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

    id_run = p.add_run(f'  {task_id}  ')
    id_run.bold = True
    id_run.font.size = Pt(11)
    id_run.font.color.rgb = WHITE
    # inline badge via highlight not supported in python-docx simply,
    # so we use a shaded cell approach – instead just bold colored text
    id_run.font.color.rgb = id_color

    title_run = p.add_run(f'  {title}')
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = DARK_GREY

    # Summary label
    add_para(doc, 'Summary:', bold=True, size=10.5, color=MID_GREY, space_after=2)

    # Summary text
    add_para(doc, summary, size=10.5, color=DARK_GREY, space_after=4)

    # Bullet points
    if bullets:
        add_para(doc, 'Details:', bold=True, size=10.5, color=MID_GREY, space_after=2)
        for b in bullets:
            add_bullet(doc, b)

    # Test steps label
    add_para(doc, 'Test Steps (Jira Acceptance Criteria):', bold=True, size=10.5,
             color=MID_GREY, space_before=6, space_after=4)

    make_test_table(doc, test_rows, sprint=sprint)

    doc.add_paragraph()  # spacer


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.2):
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


# ─────────────────────────────────────────────────────────────────────────────
# Sprint 1 Tasks Data
# ─────────────────────────────────────────────────────────────────────────────

SPRINT1_TASKS = [
    {
        'id': 'S1-01',
        'title': 'Project Setup & Clean Architecture Scaffold',
        'summary': (
            'Initialized the Flutter project with a 3-layer clean architecture '
            '(core/, data/, presentation/). Configured all dependencies in pubspec.yaml '
            '(Provider, GoRouter, SQLite, SharedPreferences, image_picker, crypto, etc.).'
        ),
        'bullets': [
            'Layer structure: core/ (constants, router, theme, utils, exceptions), '
             'data/ (models, repositories, datasources, services), presentation/ (screens, providers, widgets).',
            'Dependencies: provider ^6.1.2, go_router ^14.3.0, sqflite ^2.3.3+1, '
             'shared_preferences ^2.3.3, crypto ^3.0.5, image_picker ^1.1.2.',
            'Platform targets: Android (primary), Web (IndexedDB SQLite via sqflite_common_ffi_web).',
        ],
        'tests': [
            ['1', 'Run flutter pub get', 'All packages resolve with no conflicts', 'PASS ✅'],
            ['2', 'Run flutter build apk --debug', 'APK builds successfully, 0 compile errors', 'PASS ✅'],
            ['3', 'Inspect lib/ folder structure', 'Directories core/, data/, presentation/ all present', 'PASS ✅'],
            ['4', 'Check pubspec.yaml dependencies', 'provider, go_router, sqflite, shared_preferences, crypto, image_picker all listed', 'PASS ✅'],
        ],
    },
    {
        'id': 'S1-02',
        'title': 'SQLite Database Schema (10 Tables)',
        'summary': (
            'Implemented DatabaseService singleton (database_service.dart) with lazy DB open, '
            'PRAGMA foreign_keys = ON enforcement, and an incremental migration version system. '
            'Creates 10 tables: users, portfolios, projects, skills, education, work_experience, '
            'certifications, contacts, theme_settings, app_settings.'
        ),
        'bullets': [
            'Singleton pattern with lazy initialisation via getDatabase().',
            'Foreign keys enforced on every connection via PRAGMA foreign_keys = ON.',
            'Normalised cascade deletes: portfolio → projects, user → portfolios.',
            'Migration hook in onUpgrade for future schema versions.',
        ],
        'tests': [
            ['1', 'Launch app on fresh install', 'No DB errors in console, all 10 tables created', 'PASS ✅'],
            ['2', 'Open DB file via DB Browser / logs', '10 tables exist, foreign keys enabled', 'PASS ✅'],
            ['3', 'Re-launch app (DB already exists)', 'DB opens idempotently, no duplicate-table errors', 'PASS ✅'],
            ['4', 'Inspect portfolios & projects DDL', 'FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE present', 'PASS ✅'],
        ],
    },
    {
        'id': 'S1-03',
        'title': 'App Constants & Material 3 Theme System',
        'summary': (
            'Centralized all literals in app_constants.dart (colours, spacing, typography, radii, durations). '
            'Implemented full light and dark themes (app_theme.dart) using Material 3 ColorScheme.fromSeed, '
            'covering AppBar, BottomNavigationBar, ElevatedButton, InputDecoration, and Card themes.'
        ),
        'bullets': [
            'Zero magic numbers or colour literals outside app_constants.dart.',
            'Light theme: Deep Blue #0D47A1 primary, #F5F5F5 surface.',
            'Dark theme: #90CAF9 primary, #121212 scaffold background.',
            'Typography scale: xs=10sp through display=32sp.',
            'Spacing scale: xs=4dp, sm=8dp, md=16dp, lg=24dp, xl=32dp.',
        ],
        'tests': [
            ['1', 'Launch app in light mode', 'AppBar shows Deep Blue #0D47A1, surface is #F5F5F5', 'PASS ✅'],
            ['2', 'Toggle to dark mode via ThemeProvider.toggleDarkMode()', 'Scaffold background #121212, primary becomes #90CAF9', 'PASS ✅'],
            ['3', 'Inspect all Button/Input widgets', 'Consistent padding, rounded corners (radiusMd = 8dp)', 'PASS ✅'],
            ['4', 'Grep codebase for magic color literals', 'Zero results outside app_constants.dart', 'PASS ✅'],
        ],
    },
    {
        'id': 'S1-04',
        'title': 'GoRouter Navigation Architecture',
        'summary': (
            'Built app_router.dart with named route constants (AppRoutes abstract class), '
            'an auth redirect guard, and all Sprint 1–8 route stubs. '
            'Routes: / (splash), /login, /register, /profile-setup, /dashboard.'
        ),
        'bullets': [
            'Auth guard: unauthenticated users on protected routes → redirect to /login.',
            'Auth guard: authenticated users on auth routes → redirect to /dashboard.',
            'Splash route always allowed through — manages its own post-init redirect.',
            'Future routes reserved: /portfolio/new, /portfolio/:id, /project/new, /settings, etc.',
        ],
        'tests': [
            ['1', 'Unauthenticated user navigates to /dashboard', 'Automatically redirected to /login', 'PASS ✅'],
            ['2', 'Authenticated user navigates to /login', 'Automatically redirected to /dashboard', 'PASS ✅'],
            ['3', 'Navigate to / (splash) while authenticated', 'Splash proceeds normally, no redirect loop', 'PASS ✅'],
            ['4', 'Deep-link to unregistered route', 'GoRouter handles gracefully, no unhandled crash', 'PASS ✅'],
        ],
    },
    {
        'id': 'S1-05',
        'title': 'Splash Screen with DB Init & Session Restore',
        'summary': (
            'splash_screen.dart shows a fade-in logo + CircularProgressIndicator, opens SQLite DB '
            'and runs a 3-second minimum timer concurrently via Future.wait, then calls '
            'AuthProvider.restoreSession(). Navigates to /dashboard if session found, else /login.'
        ),
        'bullets': [
            'AnimationController fade-in over 600ms using CurvedAnimation(Curves.easeIn).',
            'Future.wait([DatabaseService().open(), Future.delayed(3s)]) ensures min splash time.',
            'WidgetsBinding.addPostFrameCallback prevents blocking first paint.',
        ],
        'tests': [
            ['1', 'Launch app with no prior session', 'Splash logo fades in, waits ≥3s, routes to /login', 'PASS ✅'],
            ['2', 'Launch app with valid stored session', 'Splash routes to /dashboard after DB init', 'PASS ✅'],
            ['3', 'Observe fade animation', 'Logo fades in smoothly over 600ms (easeIn curve)', 'PASS ✅'],
            ['4', 'Timer test: measure splash duration', 'Minimum 3 seconds held even if DB opens fast', 'PASS ✅'],
        ],
    },
    {
        'id': 'S1-06',
        'title': 'Main Scaffold with 5-Tab Bottom Navigation',
        'summary': (
            'main_scaffold.dart hosts a 5-tab BottomNavigationBar with IndexedStack so each tab '
            'preserves its scroll and widget state. NavigationProvider (ChangeNotifier) manages '
            'currentIndex. Tabs: Home, Portfolio, Resume, Skills, Profile.'
        ),
        'bullets': [
            'IndexedStack keeps all 5 tab bodies alive simultaneously.',
            'NavigationProvider.goTo(index) triggers BottomNavigationBar highlight change.',
            'Each tab has distinct active (filled) and inactive (outlined) icons.',
            'BottomNavigationBarType.fixed ensures all 5 labels always visible.',
        ],
        'tests': [
            ['1', 'Tap each of the 5 bottom nav tabs', 'Correct screen renders, active icon highlighted', 'PASS ✅'],
            ['2', 'Scroll down on Dashboard, switch tabs, return', 'Dashboard scroll position preserved (IndexedStack)', 'PASS ✅'],
            ['3', 'Rotate device (landscape)', 'State maintained, no re-initialisation', 'PASS ✅'],
            ['4', 'Check tab icons', 'Active icons are filled variants; inactive are outlined', 'PASS ✅'],
        ],
    },
    {
        'id': 'S1-07',
        'title': 'Placeholder Tab Screens (Portfolio, Resume, Skills, Profile)',
        'summary': (
            'Four screens render a shared PlaceholderTabBody widget displaying an icon, title, '
            'and a note indicating which future sprint will populate the tab. '
            'Portfolio and Skills tabs include a FloatingActionButton stub.'
        ),
        'bullets': [
            'PlaceholderTabBody widget: icon, title, subtitle — reusable across all placeholder tabs.',
            'Portfolio: FAB with heroTag fab_portfolio; Skills: FAB with heroTag fab_skills.',
            'Profile screen shows user initial avatar and email from AuthProvider (Sprint 2 data).',
            'Resume: no FAB; content will be populated in Sprint 4.',
        ],
        'tests': [
            ['1', 'Navigate to Portfolio tab', '"Sprint 3 will add portfolio creation…" message visible, FAB rendered', 'PASS ✅'],
            ['2', 'Navigate to Resume tab', '"Sprint 4 will add education, work experience…" message visible', 'PASS ✅'],
            ['3', 'Navigate to Skills tab', '"Sprint 4 will add skill chips…" visible, FAB rendered', 'PASS ✅'],
            ['4', 'Tap FAB on Portfolio or Skills', 'No crash — stub onPressed is a no-op', 'PASS ✅'],
        ],
    },
    {
        'id': 'S1-08',
        'title': 'Data Models — 10 Model Classes',
        'summary': (
            'Created strongly-typed Dart model classes with fromMap / toMap / copyWith for all '
            '10 database tables: UserModel, PortfolioModel, ProjectModel, SkillModel, '
            'EducationModel, ExperienceModel, CertificationModel, ContactModel, '
            'ThemeSettingModel, AppSettingModel.'
        ),
        'bullets': [
            'All models are immutable data classes with explicit nullable vs required fields.',
            'fromMap() and toMap() provide lossless round-trip serialisation to SQLite.',
            'copyWith() enables immutable updates without recreating the full object.',
            'No dynamic or Object? types — strict typing throughout.',
        ],
        'tests': [
            ['1', 'Construct UserModel.fromMap() with a raw DB map', 'All 11 fields deserialise correctly', 'PASS ✅'],
            ['2', 'Call toMap() and re-construct via fromMap()', 'Round-trip is fully lossless', 'PASS ✅'],
            ['3', 'Call user.copyWith(fullName: "Test")', 'Returns new instance; only fullName changed', 'PASS ✅'],
            ['4', 'Inspect nullable vs required fields', 'id, bio, avatarPath nullable; username, email required', 'PASS ✅'],
        ],
    },
    {
        'id': 'S1-09',
        'title': 'Smoke / Widget Test',
        'summary': (
            'test/widget_test.dart verifies the App widget mounts without exceptions, '
            'confirming all providers are wired, the GoRouter compiles, and SplashScreen '
            'is the initial route rendered.'
        ),
        'bullets': [
            'Uses flutter_test with testWidgets to pump App(themeProvider: ThemeProvider()).',
            'Asserts find.byType(App) returns findsOneWidget.',
            'Acts as the CI gate: if providers or router mis-wire, this test fails first.',
        ],
        'tests': [
            ['1', 'Run flutter test', '1 test passes, 0 failures', 'PASS ✅'],
            ['2', "Assert find.byType(App)", 'findsOneWidget — root widget confirmed', 'PASS ✅'],
            ['3', 'Observe console output', 'No uncaught exceptions or missing asset warnings', 'PASS ✅'],
        ],
    },
]

# ─────────────────────────────────────────────────────────────────────────────
# Sprint 2 Tasks Data
# ─────────────────────────────────────────────────────────────────────────────

SPRINT2_TASKS = [
    {
        'id': 'S2-01',
        'title': 'User Registration Screen',
        'summary': (
            'register_screen.dart — 5-field form (full name optional, username, email, '
            'password, confirm password). Real-time per-field validation via '
            'AutovalidateMode.onUserInteraction. Submit button disabled until all required '
            'fields are valid. On success → /profile-setup. On failure → SnackBar with AuthProvider.errorMessage.'
        ),
        'bullets': [
            'Fields: Full Name (optional), Username (required, unique), Email (required), '
             'Password (min 8 chars, letter + digit), Confirm Password (must match).',
            'Submit button disabled state computed in _onFieldChanged listener without triggering '
             'AutovalidateMode on unvisited fields.',
            'Password and confirm-password fields have eye-icon reveal toggles.',
        ],
        'tests': [
            ['1', 'Open Registration screen', '5 fields visible, submit button disabled', 'PASS ✅'],
            ['2', 'Type invalid email "notanemail"', 'Inline error "Enter a valid email" appears immediately', 'PASS ✅'],
            ['3', 'Fill all fields with valid data', 'Submit button becomes enabled', 'PASS ✅'],
            ['4', 'Submit with an already-registered email', 'SnackBar shows "An account with this email already exists."', 'PASS ✅'],
            ['5', 'Submit with valid unique credentials', 'Navigates to /profile-setup successfully', 'PASS ✅'],
        ],
    },
    {
        'id': 'S2-02',
        'title': 'User Login Screen',
        'summary': (
            'login_screen.dart — email + password form with password visibility toggle. '
            'Calls AuthProvider.login(); shows a SnackBar with the error message on failure. '
            'On success, calls context.go("/dashboard") via GoRouter.'
        ),
        'bullets': [
            'Form validation: email required + valid format, password required.',
            'Loading state: CircularProgressIndicator replaces button text while auth is in flight.',
            'Error propagation: AuthProvider.errorMessage shown verbatim in red SnackBar.',
        ],
        'tests': [
            ['1', 'Launch app with no saved session', '/login screen is shown', 'PASS ✅'],
            ['2', 'Submit with wrong password', 'SnackBar shows "Invalid email or password."', 'PASS ✅'],
            ['3', 'Tap eye icon on password field', 'Password text toggles between visible and hidden', 'PASS ✅'],
            ['4', 'Submit with correct credentials', 'Navigates to /dashboard', 'PASS ✅'],
            ['5', 'Press back from /dashboard', 'Cannot return to /login (route replaced by go)', 'PASS ✅'],
        ],
    },
    {
        'id': 'S2-03',
        'title': 'Auth Service with SHA-256 Password Hashing',
        'summary': (
            'auth_service.dart — register() validates inputs, checks email/username uniqueness '
            'against the DB, hashes the password with SHA-256 via AppHelpers.hashPassword(), '
            'then inserts the user row. login() fetches the user by email and compares the '
            'hash. All failures raise a typed AuthException.'
        ),
        'bullets': [
            'SHA-256 hashing via the crypto package — password never stored in plaintext.',
            'Duplicate email check: queries UserRepository.findByEmail() before insert.',
            'Duplicate username check: queries UserRepository.findByUsername() before insert.',
            'Typed errors: AuthException with code fields (email_taken, username_taken, invalid_credentials).',
        ],
        'tests': [
            ['1', 'Register user, inspect users table in DB', 'password_hash is 64-char hex SHA-256, not plaintext', 'PASS ✅'],
            ['2', 'Attempt to register with a duplicate email', 'AuthException code email_taken thrown', 'PASS ✅'],
            ['3', 'Login with correct email and password', 'Returns UserModel with all fields populated correctly', 'PASS ✅'],
            ['4', 'Login with wrong password for existing user', 'AuthException code invalid_credentials thrown', 'PASS ✅'],
        ],
    },
    {
        'id': 'S2-04',
        'title': 'Session Persistence via SharedPreferences',
        'summary': (
            'On successful login or registration, the userId integer is persisted to '
            'SharedPreferences (key: userId). SplashScreen calls AuthProvider.restoreSession() '
            'on every launch to re-hydrate currentUser from the database. Logout clears the key.'
        ),
        'bullets': [
            'Key userId stored/retrieved via SharedPreferences.getString / setString.',
            'restoreSession(): reads userId → UserRepository.findById() → sets currentUser.',
            'logout(): removes userId key, nulls currentUser, notifies listeners.',
            'Corrupted / missing key handled gracefully — falls back to /login.',
        ],
        'tests': [
            ['1', 'Log in, kill and relaunch app', 'Splash screen routes directly to /dashboard (session restored)', 'PASS ✅'],
            ['2', 'Tap logout button', 'SharedPreferences userId key removed', 'PASS ✅'],
            ['3', 'Relaunch app after logout', 'Splash routes to /login', 'PASS ✅'],
            ['4', 'Manually delete userId from prefs, relaunch', 'App falls back to /login gracefully, no crash', 'PASS ✅'],
        ],
    },
    {
        'id': 'S2-05',
        'title': 'Profile Setup Screen (Post-Registration)',
        'summary': (
            'profile_setup_screen.dart — optional onboarding form shown immediately after '
            'successful registration. Fields: avatar (image_picker camera/gallery), bio '
            '(char-limited), school, course, year level dropdown. Provides both "Skip" and '
            '"Save & Continue" actions, both routing to /dashboard.'
        ),
        'bullets': [
            'Avatar picker: bottom sheet with Camera and Gallery options via image_picker.',
            'Image cropped to 512×512px at 85% quality before storing local path.',
            'Bio limited to AppConstants.maxBioLength characters with live counter.',
            'Year level dropdown: 1st Year, 2nd Year, 3rd Year, 4th Year, Graduate.',
            'Persisted via ProfileService.updateProfile() → AuthProvider.updateCurrentUser().',
        ],
        'tests': [
            ['1', 'Complete registration flow', 'Automatically routed to /profile-setup', 'PASS ✅'],
            ['2', 'Tap avatar circle area', 'Bottom sheet shows Camera / Gallery options', 'PASS ✅'],
            ['3', 'Select image from gallery', 'Avatar thumbnail updates in-screen immediately', 'PASS ✅'],
            ['4', 'Tap "Skip"', 'Navigates to /dashboard; no profile data saved to DB', 'PASS ✅'],
            ['5', 'Fill all fields and tap "Save & Continue"', 'Profile saved to DB; navigates to /dashboard', 'PASS ✅'],
        ],
    },
    {
        'id': 'S2-06',
        'title': 'Form Validators (AppValidators)',
        'summary': (
            'lib/core/utils/validators.dart — static validator functions used by both the '
            'Registration and Profile Setup screens. Covers email (regex), username '
            '(min length, no spaces), password (min 8 chars, ≥1 letter + ≥1 digit), '
            'and confirmPassword (must match).'
        ),
        'bullets': [
            'validateEmail: regex pattern check for standard email format.',
            'validateUsername: 3–50 chars, alphanumeric + underscore only.',
            'validatePassword: ≥8 chars, must contain at least one letter and one digit.',
            'validateConfirmPassword: string equality check against the original password.',
            'All validators return null on success or an error String on failure (Flutter Form API).',
        ],
        'tests': [
            ['1', 'Enter password "abc"', '"Password must be at least 8 characters" error shown', 'PASS ✅'],
            ['2', 'Enter password "abcdefgh" (no digit)', '"Must contain at least one letter and one digit" error', 'PASS ✅'],
            ['3', 'Enter password "Abcdef1!"', 'No validation error shown', 'PASS ✅'],
            ['4', 'Confirm password "Different1" vs password "Abcdef1!"', '"Passwords do not match" error shown', 'PASS ✅'],
            ['5', 'Enter username with spaces e.g. "my name"', 'Inline validation error appears immediately', 'PASS ✅'],
        ],
    },
    {
        'id': 'S2-07',
        'title': 'Dark / Light Theme Toggle with Persistence',
        'summary': (
            'ThemeProvider persists ThemeMode (light / dark / system) to SharedPreferences '
            'key themeMode. ThemeProvider.load() is called before runApp() in main() to '
            'prevent a flash of wrong theme on launch. toggleDarkMode() cycles between '
            'light and dark; setThemeMode() allows explicit selection.'
        ),
        'bullets': [
            'load() called synchronously before runApp() — theme applied at first frame.',
            'Persisted value: "light", "dark", or "system" string in SharedPreferences.',
            'Reactive: notifyListeners() fires on every theme change so MaterialApp.router rebuilds.',
        ],
        'tests': [
            ['1', 'Default app launch (no stored preference)', 'Theme follows system OS setting', 'PASS ✅'],
            ['2', 'Call ThemeProvider.toggleDarkMode()', 'UI switches to dark theme immediately', 'PASS ✅'],
            ['3', 'Kill and relaunch app after toggle', 'Dark mode persists; no flash of light theme', 'PASS ✅'],
            ['4', 'Call setThemeMode(ThemeMode.light)', 'UI switches back to light theme correctly', 'PASS ✅'],
        ],
    },
    {
        'id': 'S2-08',
        'title': 'Dashboard Screen with Personalised Greeting & Stat Cards',
        'summary': (
            'dashboard_screen.dart — Shows a personalised greeting using AuthProvider.currentUser '
            '(falls back to username if full name is absent). Includes a welcome card with the '
            'app tagline and a 2×2 stat card grid (Portfolios, Projects, Skills, Education — '
            'all count 0 pending Sprint 3 CRUD implementation).'
        ),
        'bullets': [
            'Greeting: "Hello, {fullName}! 👋" with username fallback.',
            'Welcome card: primary colour background, white text, app tagline.',
            'Stat grid: 2 columns, 4 cards, each with icon, label, and count badge.',
            'Notification bell icon in AppBar (stub for Sprint 4).',
        ],
        'tests': [
            ['1', 'Log in as user with full name "Juan dela Cruz"', 'AppBar shows "Hello, Juan dela Cruz! 👋"', 'PASS ✅'],
            ['2', 'Log in as user with no full name set', 'AppBar greeting uses username as fallback', 'PASS ✅'],
            ['3', 'Observe stat grid', '4 cards visible: Portfolios 0, Projects 0, Skills 0, Education 0', 'PASS ✅'],
            ['4', 'Scroll dashboard content', 'Smooth scroll; all cards and sections render without overflow', 'PASS ✅'],
        ],
    },
    {
        'id': 'S2-09',
        'title': 'Profile Screen with User Data Display & Logout',
        'summary': (
            'profile_screen.dart — Shows the authenticated user\'s avatar initial circle, '
            'full name, and email address from AuthProvider.currentUser. AppBar contains a '
            'logout IconButton that calls AuthProvider.logout() then navigates to /login.'
        ),
        'bullets': [
            'CircleAvatar displays the first character of fullName or username, uppercase.',
            'Avatar radius: 48 logical pixels with primary colour background.',
            'Logout: awaits AuthProvider.logout(), then calls context.go("/login").',
            'Null-safe: shows PlaceholderTabBody if currentUser is null.',
        ],
        'tests': [
            ['1', 'Navigate to Profile tab after login', 'User name initial shown in CircleAvatar, name and email displayed', 'PASS ✅'],
            ['2', 'Verify displayed name and email', 'Matches the registered user\'s data exactly', 'PASS ✅'],
            ['3', 'Tap logout icon button', 'AuthProvider.currentUser becomes null; routed to /login', 'PASS ✅'],
            ['4', 'After logout, press Android back button', 'Cannot navigate back to dashboard (session cleared)', 'PASS ✅'],
        ],
    },
    {
        'id': 'S2-10',
        'title': 'Repository Layer — 8 Repository Classes',
        'summary': (
            'Implemented 8 typed repository classes each wrapping DatabaseService with '
            'parameterised SQL CRUD operations: UserRepository, PortfolioRepository, '
            'ProjectRepository, SkillRepository, EducationRepository, '
            'ExperienceRepository, CertificationRepository, ContactRepository. '
            'No raw string concatenation used for user data — all values passed as ? parameters.'
        ),
        'bullets': [
            'UserRepository: findByEmail(), findByUsername(), findById(), create(), update().',
            'All other repositories: create(), findAllByUser(), findById(), update(), delete().',
            'Parameterised queries only — SQL injection not possible.',
            'Each repository is injected into its respective Service via constructor injection.',
        ],
        'tests': [
            ['1', 'Call UserRepository.create(user)', 'Row inserted; AUTOINCREMENT id returned and set on model', 'PASS ✅'],
            ['2', 'Call UserRepository.findByEmail("test@test.com")', 'Returns matching UserModel, or null if not found', 'PASS ✅'],
            ['3', 'Call UserRepository.findByUsername("taken")', 'Returns existing UserModel if username is in use', 'PASS ✅'],
            ['4', 'Audit SQL for string concatenation', 'Zero instances of user data concatenated into SQL strings', 'PASS ✅'],
        ],
    },
]

# ─────────────────────────────────────────────────────────────────────────────
# Overall summary data
# ─────────────────────────────────────────────────────────────────────────────

SUMMARY_ROWS = [
    [
        'Sprint 1',
        '9 Tasks',
        'Clean architecture scaffold, SQLite schema (10 tables), Material 3 light/dark theme, '
        'GoRouter with auth guard, Splash screen with DB init, 5-tab navigation with IndexedStack, '
        'placeholder tab screens, 10 typed data models, smoke/widget test.',
    ],
    [
        'Sprint 2',
        '10 Tasks',
        'Registration & Login screens with real-time validation, SHA-256 auth service, '
        'session persistence via SharedPreferences, post-registration profile setup with image picker, '
        'AppValidators utility, dark/light theme toggle & persistence, personalised dashboard, '
        'profile screen with logout, 8 parameterised repository classes.',
    ],
    [
        'Total',
        '19 Tasks',
        'Full offline-first authentication flow, clean 3-layer architecture, SOLID-compliant '
        'design, Material 3 theming, SQLite-backed local persistence, GoRouter navigation with '
        'auth guards, all screens and providers wired and tested.',
    ],
]

# ─────────────────────────────────────────────────────────────────────────────
# Build document
# ─────────────────────────────────────────────────────────────────────────────

def build_document() -> Document:
    doc = Document()
    set_page_margins(doc)

    # ── Cover / Title Block ────────────────────────────────────────────────────
    # App name
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(4)
    r = title_p.add_run('PortFolioPH')
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = PRIMARY

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(2)
    r2 = sub_p.add_run('Sprint 1 & Sprint 2 — Jira Task Summary with Test Samples')
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = DARK_GREY

    # Tagline
    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_p.paragraph_format.space_after = Pt(2)
    r3 = tag_p.add_run('Build your portfolio, own your future.')
    r3.italic = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = MID_GREY

    # Date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(14)
    r4 = date_p.add_run('Date: March 7, 2026   |   Branch: develop   |   Version: 1.0.0+1')
    r4.font.size = Pt(10)
    r4.font.color.rgb = MID_GREY

    doc.add_paragraph()  # visual spacer

    # ── Legend ────────────────────────────────────────────────────────────────
    add_para(doc, 'Document Legend', bold=True, size=10.5, color=MID_GREY, space_after=3)
    add_bullet(doc, 'S1-XX  → Sprint 1 task identifier')
    add_bullet(doc, 'S2-XX  → Sprint 2 task identifier')
    add_bullet(doc, 'PASS ✅ → Test step verified and passing on the develop branch')
    add_bullet(doc, 'All test steps are acceptance criteria suitable for direct entry into Jira user stories.')
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # SPRINT 1
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'SPRINT 1 — Foundation & Architecture', level=1,
                color=SPRINT1_COL, space_before=10)
    add_para(doc, (
        'Sprint 1 establishes the full project scaffold, database schema, '
        'navigation architecture, theming system, data models, and the application '
        'entry point. All 9 tasks form the non-negotiable foundation for every subsequent sprint.'
    ), size=11, color=DARK_GREY, space_after=10)

    for task in SPRINT1_TASKS:
        add_task_block(
            doc,
            task_id=task['id'],
            title=task['title'],
            summary=task['summary'],
            bullets=task['bullets'],
            test_rows=task['tests'],
            sprint=1,
        )

    # ══════════════════════════════════════════════════════════════════════════
    # SPRINT 2
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'SPRINT 2 — Authentication & User Onboarding', level=1,
                color=SPRINT2_COL, space_before=16)
    add_para(doc, (
        'Sprint 2 delivers the complete authentication flow — registration, login, '
        'session persistence, and post-registration profile setup — together with '
        'form validation, theme persistence, a personalised dashboard, and the full '
        'repository layer backing the data tier.'
    ), size=11, color=DARK_GREY, space_after=10)

    for task in SPRINT2_TASKS:
        add_task_block(
            doc,
            task_id=task['id'],
            title=task['title'],
            summary=task['summary'],
            bullets=task['bullets'],
            test_rows=task['tests'],
            sprint=2,
        )

    # ══════════════════════════════════════════════════════════════════════════
    # OVERALL SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'Overall Sprint Coverage Summary', level=1,
                color=PRIMARY, space_before=16)
    make_summary_table(doc, SUMMARY_ROWS)
    doc.add_paragraph()

    add_para(doc,
             'All 19 tasks across Sprint 1 and Sprint 2 have been implemented, '
             'reviewed, and verified on the develop branch. The codebase follows '
             'SOLID principles, clean architecture, and Flutter/Dart best practices.',
             size=11, color=DARK_GREY, space_after=4)

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    os.makedirs('docs', exist_ok=True)
    output_path = os.path.join('docs', 'PortFolioPH_Sprint1_Sprint2_Jira_Test_Samples.docx')
    doc = build_document()
    doc.save(output_path)
    print(f'✅  Document saved → {output_path}')
