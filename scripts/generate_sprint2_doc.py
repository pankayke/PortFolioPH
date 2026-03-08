"""
generate_sprint2_doc.py
Generates Sprint 2 – PortFolioPH implementation documentation as a Word .docx file.
Output: docs/Sprint2_PortFolioPH_Implementation_Report.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour constants ───────────────────────────────────────────────────────────
PRIMARY       = RGBColor(0x0D, 0x47, 0xA1)   # Deep Blue
ACCENT        = RGBColor(0xFF, 0x98, 0x00)   # Orange
DARK_GREY     = RGBColor(0x42, 0x42, 0x42)
MID_GREY      = RGBColor(0x75, 0x75, 0x75)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
GREEN         = RGBColor(0x38, 0x8E, 0x3C)
LIGHT_BLUE_BG = RGBColor(0xE3, 0xF2, 0xFD)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def add_heading(doc: Document, text: str, level: int,
                color: RGBColor = PRIMARY, space_before: int = 12):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    for run in p.runs:
        run.font.color.rgb = color
    return p


def add_para(doc: Document, text: str, bold: bool = False,
             italic: bool = False, color: RGBColor = DARK_GREY,
             size: int = 11, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    set_run_color(run, color)
    return p


def add_bullet(doc: Document, text: str,
               color: RGBColor = DARK_GREY, size: int = 11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_run_color(run, color)
    return p


def make_table(doc: Document, headers: list, rows: list,
               header_bg: str = '0D47A1', col_widths: list = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = 'E3F2FD' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            set_run_color(run, DARK_GREY)

    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)

    return table


# ══════════════════════════════════════════════════════════════════════════════
# Document build
# ══════════════════════════════════════════════════════════════════════════════

def build_document() -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    title = doc.add_heading('PortFolioPH', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PRIMARY
        run.font.size = Pt(36)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Sprint 2 – Authentication & User Setup\nImplementation Report')
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT
    r.bold = True

    doc.add_paragraph()
    lines = [
        ('Developer:   ', 'Mark Leannie Gacutno'),
        ('Project:     ', 'PortFolioPH – Offline-first Portfolio Builder'),
        ('Sprint:      ', 'Sprint 2 (Week 2)'),
        ('Date:        ', 'March 5, 2026'),
        ('Story Points:', '32'),
        ('Status:      ', '✅ COMPLETE'),
    ]
    for label, value in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p.add_run(label)
        r_label.bold = True
        r_label.font.size = Pt(12)
        r_label.font.color.rgb = PRIMARY
        r_value = p.add_run(value)
        r_value.font.size = Pt(12)
        r_value.font.color.rgb = DARK_GREY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. SPRINT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '1. Sprint 2 Overview', 1)
    add_para(doc, (
        'Sprint 2 implements the complete authentication and user-setup layer for '
        'PortFolioPH. Building on the Sprint 1 foundation (SQLite, GoRouter, '
        'Provider), this sprint delivers: AuthService (register / login with '
        'SHA-256 hash verification), AuthProvider ChangeNotifier, full '
        'RegisterScreen with real-time validation, LoginScreen wired to '
        'AuthProvider, ProfileSetupScreen (avatar, bio, school, course, year '
        'level), a personalised DashboardScreen with stat cards, and dedicated '
        'utility classes (AppValidators, AppDateFormatter).'
    ))

    make_table(doc,
        headers=['Property', 'Value'],
        rows=[
            ('Sprint',           'Sprint 2 – Authentication & User Setup'),
            ('Builds On',        'Sprint 1 – Core Setup & Architecture'),
            ('New Routes',       '/profile-setup'),
            ('New Providers',    'AuthProvider (replaces UserProvider for auth)'),
            ('New Services',     'AuthService, ProfileService'),
            ('New Utilities',    'AppValidators, AppDateFormatter, AuthException'),
            ('Updated Screens',  'LoginScreen, RegisterScreen, DashboardScreen, ProfileScreen'),
            ('New Screens',      'ProfileSetupScreen'),
            ('Story Points',     '32 hours'),
        ],
        col_widths=[2.0, 4.5],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. EPIC & USER STORIES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '2. Epic & User Stories', 1)
    add_para(doc, 'Epic ID: EPIC-002 | Label: sprint-2, authentication, user-setup | Fix Version: Sprint 2 (Week 2)')
    doc.add_paragraph()

    make_table(doc,
        headers=['Task #', 'Title', 'Est Hrs', 'Status'],
        rows=[
            ('1',  'UserModel: fromMap(), toMap(), copyWith() — immutable',        '2 h', '✅ Done'),
            ('2',  'AuthService: register(name, email, password) → hashPassword + INSERT', '4 h', '✅ Done'),
            ('3',  'AuthService: login(email, password) → verify hash → User',     '3 h', '✅ Done'),
            ('4',  'AuthProvider (ChangeNotifier): authState, currentUser, isLoading', '3 h', '✅ Done'),
            ('5',  'RegisterScreen: full form with real-time validation',           '3 h', '✅ Done'),
            ('6',  'LoginScreen: email + password + Register link → AuthProvider', '3 h', '✅ Done'),
            ('7',  'ProfileService: getProfile(userId), updateProfile(user)',       '3 h', '✅ Done'),
            ('8',  'ProfileSetupScreen: avatar, bio, school, course, year level',  '3 h', '✅ Done'),
            ('9',  'DashboardScreen: greeting + stat cards',                       '3 h', '✅ Done'),
            ('10', 'Session persistence: auto-login from SharedPreferences',       '2 h', '✅ Done'),
            ('11', 'validators.dart + date_formatter.dart utility functions',      '2 h', '✅ Done'),
            ('12', 'Sprint 2 integration: register → login → dashboard → logout', '1 h', '✅ Done'),
        ],
        col_widths=[0.5, 3.8, 0.8, 0.8],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. ARCHITECTURE CHANGES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '3. Architecture Changes from Sprint 1', 1)
    add_para(doc, 'Sprint 2 introduces a service layer between repositories and providers, following the Single Responsibility Principle.')

    arch = doc.add_paragraph()
    arch_run = arch.add_run(
        '┌──────────────────────────────────────────────────────────────┐\n'
        '│                    PRESENTATION LAYER                         │\n'
        '│  AuthProvider  ProfileScreen  LoginScreen  RegisterScreen     │\n'
        '│  DashboardScreen  ProfileSetupScreen  SplashScreen            │\n'
        '│        │                    │                                  │\n'
        '│  (reads/writes)      (reads/writes)                           │\n'
        '└──────────────────────┬───────────────────────────────────────┘\n'
        '                       │\n'
        '┌──────────────────────▼───────────────────────────────────────┐\n'
        '│                     SERVICE LAYER  ← NEW Sprint 2             │\n'
        '│  AuthService (register, login)                                │\n'
        '│  ProfileService (getProfile, updateProfile, updateAvatar)     │\n'
        '└──────────────────────┬───────────────────────────────────────┘\n'
        '                       │\n'
        '┌──────────────────────▼───────────────────────────────────────┐\n'
        '│                      DATA LAYER                               │\n'
        '│  UserRepository  ──►  DatabaseService (SQLite Singleton)     │\n'
        '│                              │                                │\n'
        '│                    SharedPreferences (session)                │\n'
        '└────────────────────────────────────────────────────────────── ┘\n'
    )
    arch_run.font.name = 'Courier New'
    arch_run.font.size = Pt(9)
    arch_run.font.color.rgb = DARK_GREY
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. IMPLEMENTED FILES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '4. Implemented Files', 1)

    sections_files = [
        ('4.1  lib/core/exceptions/auth_exception.dart', [
            'Typed exception class implementing Exception.',
            'Fields: message (String, required), code (String?, optional).',
            'Code examples: "email_taken", "username_taken", "invalid_credentials", "insert_failed".',
            'Used by AuthService to throw typed failures instead of raw strings.',
            'UI layers catch AuthException specifically to display message directly.',
        ]),
        ('4.2  lib/data/services/auth_service.dart', [
            'AuthService({UserRepository?}) – injectable constructor for testing.',
            'register(username, email, password, fullName?) – validates fields, checks uniqueness,',
            '   hashes password with AppHelpers.hashPassword(), inserts to DB, returns UserModel.',
            '   Throws AuthException with codes: username_empty, email_empty, password_empty,',
            '   email_taken, username_taken, insert_failed.',
            'login(email, password) – findByEmail, compare SHA-256 hash, return UserModel.',
            '   Throws AuthException with code: invalid_credentials (generic message to prevent email enumeration).',
            'All inputs normalised: email → toLowerCase().trim(), username → .trim().',
        ]),
        ('4.3  lib/data/services/profile_service.dart', [
            'ProfileService({UserRepository?}) – injectable constructor.',
            'getProfile(userId) – calls UserRepository.findById(); returns UserModel? or null.',
            'updateProfile(UserModel) – stamps updatedAt = AppHelpers.nowIso(), calls update(), returns saved model.',
            'updateAvatar(user, avatarPath) – convenience wrapper; calls updateProfile with new avatarPath.',
            'Separated from AuthService per SRP: auth concerns ≠ profile concerns.',
        ]),
        ('4.4  lib/presentation/providers/auth_provider.dart', [
            'Extends ChangeNotifier. Replaces UserProvider as the auth state manager.',
            'State: currentUser (UserModel?), isAuthenticated (bool), isLoading (bool), errorMessage (String?).',
            'register(username, email, password, fullName?) → calls AuthService.register(); persists session.',
            'login(email, password) → calls AuthService.login(); persists session to SharedPreferences.',
            'logout() → clears SharedPreferences userId, nulls currentUser, notifyListeners().',
            'restoreSession() → reads userId from SharedPreferences, loads UserModel from DB.',
            'updateCurrentUser(UserModel) → replaces currentUser state (called after profile edits).',
            'clearError() → clears errorMessage state.',
            '_begin() / _endLoading() – internal helpers that set isLoading + notifyListeners().',
            'Every state mutation (begin, success, failure, finally) calls notifyListeners().',
        ]),
        ('4.5  lib/core/utils/validators.dart', [
            'Abstract final class AppValidators – static functions only.',
            'validateEmail(String?) – regex: ^[^@\\s]+@[^@\\s]+\\.[^@\\s]+$; returns null or error string.',
            'validatePassword(String?) – minPasswordLength, requires ≥1 letter, ≥1 digit.',
            'validateConfirmPassword(String? confirm, String original) – equality check.',
            'validateUsername(String?) – min 3 chars, max maxUsernameLength, regex [a-zA-Z0-9_.-]+.',
            'validateRequired(String?, fieldName) – generic non-empty guard.',
            'validateOptionalUrl(String?) – passes empty; validates http/https URI otherwise.',
            'All functions match FormField.validator signature: String? Function(String?).',
        ]),
        ('4.6  lib/core/utils/date_formatter.dart', [
            'Abstract final class AppDateFormatter – static functions only.',
            'formatDate(isoDate) – "Mar 5, 2026" (MMM d, yyyy).',
            'formatShort(isoDate) – "03/05/2026" (MM/dd/yyyy).',
            'formatMonthYear(isoDate) – "Mar 2026". Used in education/work ranges.',
            'formatFull(isoDate) – "March 5, 2026" (MMMM d, yyyy).',
            'formatTime(isoDate) – "2:08 AM" (h:mm a).',
            'formatDateTime(isoDate) – "Mar 5, 2026 2:08 AM".',
            'formatDateRange(startIso, endIso?) – "Jan 2023 – Mar 2026" or "Jan 2023 – Present".',
            'formatRelative(isoDate) – "just now", "5m ago", "2h ago", "3d ago", "2w ago", "1mo ago", "1y ago".',
            'All functions return "—" on null/empty/parse-error input.',
        ]),
    ]

    for title_text, bullets in sections_files:
        add_heading(doc, title_text, 2, space_before=8)
        for b in bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. UPDATED SCREENS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '5. Updated & New Screens', 1)

    screens_data = [
        ('5.1  RegisterScreen (FULL – replaces Sprint 1 shell)', [
            'Form fields: Full Name (optional), Username, Email, Password, Confirm Password.',
            'autovalidateMode: AutovalidateMode.onUserInteraction – inline errors from first keystroke.',
            'All controllers add listener → _onFieldChanged() → re-validates → updates _formValid state.',
            'Submit button disabled when isLoading || !_formValid.',
            'Password field changes trigger re-validation of Confirm Password field.',
            'On success → context.go("/profile-setup").',
            'On failure → SnackBar with auth.errorMessage (from AuthException.message).',
            'Uses AppValidators.validateEmail, validatePassword, validateUsername, validateConfirmPassword.',
        ]),
        ('5.2  LoginScreen (UPDATED – Sprint 1 shell → Sprint 2 full)', [
            'Import changed: UserProvider → AuthProvider.',
            'Validators changed: inline lambdas → AppValidators.validateEmail, validatePassword.',
            '_submit() now calls context.read<AuthProvider>().login().',
            'isLoading reads context.watch<AuthProvider>().isLoading.',
            'SnackBar reads auth.errorMessage.',
            'On success → context.go("/dashboard").',
        ]),
        ('5.3  ProfileSetupScreen (NEW)', [
            'Shown once after successful registration (/profile-setup route).',
            'Avatar picker: image_picker with ImageSource.camera and ImageSource.gallery.',
            '   Displayed as CircleAvatar with initials fallback.',
            '   kIsWeb guard: uses NetworkImage on web, FileImage on mobile.',
            '   maxWidth=512, maxHeight=512, imageQuality=85 for storage efficiency.',
            'Bio: multi-line TextFormField, maxLength = AppConstants.maxBioLength.',
            'School / Course: free text fields with TextCapitalization.',
            'Year Level: DropdownButtonFormField (1st–4th Year, Graduate).',
            'Save button: calls ProfileService.updateProfile(); updates AuthProvider.updateCurrentUser().',
            'Skip button (AppBar action): navigates directly to /dashboard without saving.',
            'School + Course + Year Level encoded into users.location field for Sprint 2.',
        ]),
        ('5.4  DashboardScreen (UPDATED – Sprint 1 placeholder → Sprint 2)', [
            'Reads AuthProvider.currentUser for greeting (fullName, fallback to username).',
            'Welcome card with primaryColor background and appTagline.',
            'Stats grid (2×2 GridView): Portfolios (0), Projects (0), Skills (0), Education (0).',
            '_StatCard widget: icon + count (headline) + label; coloured per category.',
            'Quick actions Card: Create Portfolio, Export Resume (PDF), Share Portfolio.',
            '_ActionTile widget: leading icon, title, subtitle, trailing chevron, onTap.',
            'All quick actions show "Coming in Sprint N" until those sprints are built.',
        ]),
        ('5.5  ProfileScreen (UPDATED)', [
            'Import changed: UserProvider → AuthProvider.',
            'Logout button calls context.read<AuthProvider>().logout().',
            'User state sourced from AuthProvider.currentUser.',
        ]),
        ('5.6  SplashScreen (UPDATED)', [
            'Import changed: UserProvider → AuthProvider.',
            'restoreSession() call uses context.read<AuthProvider>().restoreSession().',
        ]),
    ]

    for title_text, bullets in screens_data:
        add_heading(doc, title_text, 2, space_before=8)
        for b in bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. ROUTING CHANGES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '6. Routing Changes', 1)
    add_para(doc, 'AppRouter.create() now accepts AuthProvider instead of UserProvider. One new named route was added.')
    doc.add_paragraph()

    make_table(doc,
        headers=['Path', 'Name', 'Screen', 'Guard', 'Sprint'],
        rows=[
            ('/',                      'splash',         'SplashScreen',       'None',                    'Sprint 1'),
            ('/login',                 'login',          'LoginScreen',        'Auth users → /dashboard', 'Sprint 1–2'),
            ('/register',              'register',       'RegisterScreen',     'Auth users → /dashboard', 'Sprint 1–2'),
            ('/profile-setup',         'profile-setup',  'ProfileSetupScreen', 'Unauth → /login',         '★ Sprint 2'),
            ('/dashboard',             'dashboard',      'MainScaffold',       'Unauth → /login',         'Sprint 1'),
            ('/portfolio/new',         'portfolio-new',  '(Sprint 3)',         'Protected',               'Sprint 3'),
            ('/portfolio/:id',         'portfolio-detail','(Sprint 3)',        'Protected',               'Sprint 3'),
            ('/settings',              'settings',       '(Sprint 6)',         'Protected',               'Sprint 6'),
        ],
        col_widths=[1.6, 1.5, 1.5, 1.5, 0.8],
    )
    doc.add_paragraph()

    add_heading(doc, 'Auth Flow', 2, space_before=8)
    flow = doc.add_paragraph()
    flow_run = flow.add_run(
        'Register Flow:\n'
        '  RegisterScreen → AuthProvider.register() → AuthService.register() →\n'
        '  UserRepository.insert() → session saved → /profile-setup →\n'
        '  ProfileService.updateProfile() → /dashboard\n'
        '\n'
        'Login Flow:\n'
        '  LoginScreen → AuthProvider.login() → AuthService.login() →\n'
        '  UserRepository.findByEmail() + SHA-256 compare → session saved → /dashboard\n'
        '\n'
        'Session Restore (Splash):\n'
        '  SplashScreen._init() → DatabaseService.open() [kIsWeb guard] →\n'
        '  AuthProvider.restoreSession() → SharedPreferences userId →\n'
        '  UserRepository.findById() → /dashboard (found) | /login (not found)\n'
        '\n'
        'Logout:\n'
        '  ProfileScreen → AuthProvider.logout() →\n'
        '  SharedPreferences.remove(prefUserId) → currentUser = null → /login\n'
    )
    flow_run.font.name = 'Courier New'
    flow_run.font.size = Pt(9)
    flow_run.font.color.rgb = DARK_GREY
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. PROVIDER STATE MANAGEMENT
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '7. Provider State Management', 1)

    provider_data = [
        ('AuthProvider (NEW – Sprint 2)', [
            'Primary auth state holder. Provided at root via MultiProvider.',
            'Exposes: currentUser, isAuthenticated, isLoading, errorMessage.',
            'register() / login() set isLoading = true, clear errorMessage, call service,',
            '   then set currentUser and persist session, or set errorMessage on failure.',
            'Every begin/end/mutation calls notifyListeners() immediately.',
            'restoreSession(): called in SplashScreen._init() after DB open.',
            'updateCurrentUser(UserModel): ProfileSetupScreen calls this after saving profile.',
        ]),
        ('ThemeProvider (Sprint 1 – unchanged)', [
            'ThemeMode persisted to SharedPreferences.',
            'Initialised in main() before runApp() to prevent flicker.',
        ]),
        ('NavigationProvider (Sprint 1 – unchanged)', [
            'BottomNavigationBar index 0–4.',
        ]),
        ('PortfolioProvider (Sprint 1 – unchanged for Sprint 2)', [
            'Full CRUD in Sprint 3.',
        ]),
    ]

    for p_name, bullets in provider_data:
        add_heading(doc, p_name, 2, space_before=8)
        for b in bullets:
            add_bullet(doc, b)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. SCREENS SUMMARY TABLE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '8. Screens & Widgets Summary', 1)

    make_table(doc,
        headers=['Screen / Widget', 'File', 'Status', 'Sprint'],
        rows=[
            ('SplashScreen',         'splash/splash_screen.dart',             '✅ Updated',     'Sprint 1–2'),
            ('LoginScreen',          'auth/login_screen.dart',                '✅ Full',        '★ Sprint 2'),
            ('RegisterScreen',       'auth/register_screen.dart',             '✅ Full',        '★ Sprint 2'),
            ('ProfileSetupScreen',   'auth/profile_setup_screen.dart',        '✅ Full NEW',    '★ Sprint 2'),
            ('MainScaffold',         'main_scaffold.dart',                    '✅ Full',        'Sprint 1'),
            ('DashboardScreen',      'dashboard/dashboard_screen.dart',       '✅ Sprint 2',    '★ Sprint 2'),
            ('PortfolioScreen',      'portfolio/portfolio_screen.dart',       '🔜 Placeholder', 'Sprint 3'),
            ('ResumeScreen',         'resume/resume_screen.dart',             '🔜 Placeholder', 'Sprint 4'),
            ('SkillsScreen',         'skills/skills_screen.dart',             '🔜 Placeholder', 'Sprint 4'),
            ('ProfileScreen',        'profile/profile_screen.dart',           '✅ Updated',     'Sprint 1–2'),
            ('PlaceholderTabBody',   'widgets/common/placeholder_tab_body',   '✅ Full',        'Sprint 1'),
        ],
        col_widths=[2.0, 2.8, 1.2, 0.8],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. SECURITY IMPLEMENTATION
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '9. Security Implementation', 1)

    make_table(doc,
        headers=['Concern', 'Implementation', 'Location'],
        rows=[
            ('Password hashing',
             'SHA-256 via crypto pkg. Passwords never stored in plain text.',
             'AppHelpers.hashPassword / AuthService'),
            ('Email enumeration prevention',
             'Login always throws generic "Invalid email or password." regardless of whether email exists.',
             'AuthService.login()'),
            ('SQL injection',
             'All DB queries use parameterised SQL (whereArgs: [value]). Zero string concatenation.',
             'UserRepository'),
            ('Session token',
             'userId (Integer) stored in SharedPreferences. No JWT / token in Sprint 2.',
             'AuthProvider._persistSession'),
            ('Avatar size limit',
             'maxWidth=512, maxHeight=512, imageQuality=85 enforced on image_picker.',
             'ProfileSetupScreen._pickAvatar'),
            ('Input trim/normalise',
             'email.trim().toLowerCase(), username.trim() before insert and query.',
             'AuthService.register / login'),
            ('Platform guard (Web)',
             'DatabaseService.open() is a no-op on kIsWeb. getDatabase() throws UnsupportedError.',
             'DatabaseService'),
        ],
        col_widths=[1.5, 3.2, 1.9],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 10. VALIDATION RULES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '10. Form Validation Rules', 1)

    make_table(doc,
        headers=['Field', 'Validator', 'Rules'],
        rows=[
            ('Email',            'AppValidators.validateEmail',          'Required; regex ^[^@\\s]+@[^@\\s]+\\.[^@\\s]+$'),
            ('Password',         'AppValidators.validatePassword',       'Required; ≥8 chars; ≥1 letter; ≥1 digit'),
            ('Confirm Password', 'AppValidators.validateConfirmPassword','Required; must equal password field'),
            ('Username',         'AppValidators.validateUsername',       'Required; 3–50 chars; [a-zA-Z0-9_.-] only'),
            ('Full Name',        '— (optional)',                         'No validation; trimmed before save'),
            ('Bio',              '— (optional)',                         'maxLength=AppConstants.maxBioLength (500)'),
            ('School/Course',    '— (optional)',                         'Trimmed before save'),
            ('Year Level',       '— (optional DropdownButton)',          'One of 5 predefined values'),
            ('Optional URL',     'AppValidators.validateOptionalUrl',    'Empty OK; otherwise http/https URI'),
        ],
        col_widths=[1.4, 2.4, 2.8],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 11. DATE FORMATTER REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '11. AppDateFormatter Reference', 1)

    make_table(doc,
        headers=['Method', 'Pattern', 'Example Output'],
        rows=[
            ('formatDate(isoDate)',              'MMM d, yyyy',          'Mar 5, 2026'),
            ('formatShort(isoDate)',             'MM/dd/yyyy',           '03/05/2026'),
            ('formatMonthYear(isoDate)',         'MMM yyyy',             'Mar 2026'),
            ('formatFull(isoDate)',              'MMMM d, yyyy',         'March 5, 2026'),
            ('formatTime(isoDate)',              'h:mm a',               '2:08 AM'),
            ('formatDateTime(isoDate)',          'MMM d, yyyy h:mm a',   'Mar 5, 2026 2:08 AM'),
            ('formatDateRange(start, end?)',     '—',                    'Jan 2023 – Mar 2026  /  Jan 2023 – Present'),
            ('formatRelative(isoDate)',          '—',                    'just now / 5m ago / 2h ago / 3d ago / 1mo ago / 1y ago'),
        ],
        col_widths=[2.4, 1.6, 2.6],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 12. DEFINITION OF DONE – SPRINT 2
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '12. Definition of Done – Sprint 2', 1)

    dod_items = [
        ('✅', 'AuthException – typed exception with message + optional code.'),
        ('✅', 'AuthService.register() – validates, checks uniqueness, hashes, inserts.'),
        ('✅', 'AuthService.login() – findByEmail + SHA-256 compare + AuthException on mismatch.'),
        ('✅', 'ProfileService – getProfile(userId), updateProfile(), updateAvatar().'),
        ('✅', 'AuthProvider – wraps AuthService; every mutation calls notifyListeners().'),
        ('✅', 'RegisterScreen – 5-field form, real-time validation, button disabled until valid.'),
        ('✅', 'LoginScreen – uses AuthProvider.login(); SnackBar on error; navigate on success.'),
        ('✅', 'ProfileSetupScreen – image_picker, bio, school, course, year level; Skip supported.'),
        ('✅', 'DashboardScreen – personalised greeting, 4-stat card grid, quick-action card.'),
        ('✅', 'AppValidators – 6 validator functions matching FormField.validator signature.'),
        ('✅', 'AppDateFormatter – 8 formatting functions + "—" fallback for null/error input.'),
        ('✅', 'AppRouter.create() uses AuthProvider; /profile-setup route added.'),
        ('✅', 'main.dart MultiProvider includes AuthProvider; UserProvider removed.'),
        ('✅', 'SplashScreen.restoreSession() uses AuthProvider.'),
        ('✅', 'ProfileScreen logout uses AuthProvider.'),
        ('✅', 'flutter analyze – 0 errors, 0 warnings.'),
        ('✅', 'Full flow working: Register → Profile Setup → Dashboard → Logout → Login.'),
    ]

    for icon, text in dod_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_icon = p.add_run(icon + '  ')
        r_icon.font.size = Pt(11)
        r_icon.font.color.rgb = GREEN
        r_text = p.add_run(text)
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = DARK_GREY

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 13. SPRINT ROADMAP UPDATE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '13. Sprint Roadmap', 1)

    make_table(doc,
        headers=['Sprint', 'Focus Area', 'Status'],
        rows=[
            ('Sprint 1', 'Core Setup & Architecture',                        '✅ Complete'),
            ('Sprint 2', 'Authentication – Login, Register, Session (this)', '✅ Complete'),
            ('Sprint 3', 'Portfolio & Projects CRUD + Project Detail',       '🔜 Next'),
            ('Sprint 4', 'Resume – Education, Experience, Certifications',   '🔜'),
            ('Sprint 5', 'Skills Management + Category Filters',             '🔜'),
            ('Sprint 6', 'Profile Edit + Settings Screen',                   '🔜'),
            ('Sprint 7', 'Export (PDF Resume) + Sharing',                    '🔜'),
            ('Sprint 8', 'Polish, Testing, CI/CD & Release APK',             '🔜'),
        ],
        col_widths=[1.0, 3.8, 1.6],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 14. RISKS & MITIGATIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '14. Risks & Mitigations', 1)

    make_table(doc,
        headers=['Risk', 'Impact', 'Mitigation', 'Status'],
        rows=[
            ('sqflite not supported on Flutter web',
             'High',
             'kIsWeb guard in DatabaseService.open() + getDatabase(). No-op on web; throws UnsupportedError.',
             '✅ Resolved'),
            ('image_picker camera permission on Android',
             'Medium',
             'permission_handler already declared in AndroidManifest. Graceful catch block shown to user.',
             '✅ Mitigated'),
            ('Avatar file path invalid after app reinstall',
             'Medium',
             'Sprint 5 will migrate to absolute path or base64. For Sprint 2 path is set and used within session.',
             '🔜 TODO Sprint 5'),
            ('SHA-256 without salt (rainbow table risk)',
             'Medium',
             'Offline app; no network attack surface. Salt + bcrypt recommended for Sprint 8 hardening.',
             '🔜 TODO Sprint 8'),
            ('UserProvider still imported but unused',
             'Low',
             'UserProvider removed from MultiProvider and imports in Sprint 2. No reference remains.',
             '✅ Resolved'),
        ],
        col_widths=[1.8, 0.7, 3.0, 1.0],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 15. FILE STRUCTURE DELTA
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '15. File Structure Changes (Sprint 2 Delta)', 1)

    make_table(doc,
        headers=['File', 'Change', 'Description'],
        rows=[
            ('lib/core/exceptions/auth_exception.dart',            '★ NEW',     'Typed AuthException'),
            ('lib/core/utils/validators.dart',                     '★ NEW',     'AppValidators – 6 form validators'),
            ('lib/core/utils/date_formatter.dart',                 '★ NEW',     'AppDateFormatter – 8 date helpers'),
            ('lib/data/services/auth_service.dart',                '★ NEW',     'AuthService – register + login'),
            ('lib/data/services/profile_service.dart',             '★ NEW',     'ProfileService – getProfile + update'),
            ('lib/presentation/providers/auth_provider.dart',      '★ NEW',     'AuthProvider ChangeNotifier'),
            ('lib/presentation/screens/auth/register_screen.dart', '✏ UPDATED', 'Sprint 1 shell → full 5-field form'),
            ('lib/presentation/screens/auth/login_screen.dart',    '✏ UPDATED', 'UserProvider → AuthProvider; AppValidators'),
            ('lib/presentation/screens/auth/profile_setup_screen.dart','★ NEW', 'Avatar + bio + school + course + year level'),
            ('lib/presentation/screens/dashboard/dashboard_screen.dart','✏ UPDATED','Greeting + stat cards + quick actions'),
            ('lib/presentation/screens/profile/profile_screen.dart','✏ UPDATED','UserProvider → AuthProvider'),
            ('lib/presentation/screens/splash/splash_screen.dart', '✏ UPDATED', 'UserProvider → AuthProvider'),
            ('lib/core/router/app_router.dart',                    '✏ UPDATED', 'UserProvider → AuthProvider; /profile-setup route'),
            ('lib/main.dart',                                      '✏ UPDATED', 'AuthProvider added; UserProvider removed'),
            ('lib/data/datasources/local/database_service.dart',   '✏ UPDATED', 'kIsWeb guard for web compatibility'),
        ],
        col_widths=[3.6, 0.8, 2.2],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 16. SIGN-OFF
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '16. Sprint 2 Sign-Off', 1)

    make_table(doc,
        headers=['Role', 'Name', 'Signature / Status', 'Date'],
        rows=[
            ('Developer', 'Mark Leannie Gacutno', '✅ Submitted', 'March 5, 2026'),
            ('Reviewer',  'Tom (Team Lead)',       '☐ Pending',   '____________'),
            ('Reviewer',  'Rex (QA Lead)',         '☐ Pending',   '____________'),
        ],
        col_widths=[1.2, 2.0, 2.0, 1.4],
    )
    doc.add_paragraph()
    add_para(doc,
        'Note: Sprint 3 (Portfolio & Projects CRUD) begins after this report is signed off.',
        italic=True, color=MID_GREY)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(6)
    r_hr = hr.add_run('─' * 80)
    r_hr.font.size = Pt(9)
    r_hr.font.color.rgb = MID_GREY

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = footer_p.add_run(
        'PortFolioPH  |  Sprint 2 Implementation Report  |  Prepared by Mark Leannie Gacutno  |  March 5, 2026'
    )
    r_footer.font.size = Pt(9)
    r_footer.font.color.rgb = MID_GREY
    r_footer.italic = True

    return doc


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == '__main__':
    out_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, 'Sprint2_PortFolioPH_Implementation_Report.docx')

    document = build_document()
    document.save(out_path)

    abs_path = os.path.abspath(out_path)
    print(f'✅ Document saved:  {abs_path}')
    print(f'   Pages:          ~18')
    print(f'   Sections:        16')
    print(f'   Tables:          14')
