"""
generate_combined_doc.py
Generates a single combined Sprint 1 + Sprint 2 Implementation Report for
PortFolioPH, covering all tasks with developer discussion notes and Jira-ready
screenshot placeholder sections.

Output: docs/PortFolioPH_Sprint1_Sprint2_Combined_Report.docx

Run:  python scripts/generate_combined_doc.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────────
PRIMARY       = RGBColor(0x0D, 0x47, 0xA1)   # Deep Blue
ACCENT        = RGBColor(0xFF, 0x98, 0x00)   # Orange
DARK_GREY     = RGBColor(0x42, 0x42, 0x42)
MID_GREY      = RGBColor(0x75, 0x75, 0x75)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
GREEN         = RGBColor(0x2E, 0x7D, 0x32)
RED           = RGBColor(0xC6, 0x28, 0x28)
ORANGE        = RGBColor(0xE6, 0x51, 0x00)
PURPLE        = RGBColor(0x6A, 0x1B, 0x9A)
TEAL          = RGBColor(0x00, 0x69, 0x6E)
LIGHT_BLUE    = RGBColor(0xE3, 0xF2, 0xFD)


# ── Low-level helpers ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level, color=PRIMARY, space_before=14):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = color
    return p


def add_para(doc, text, bold=False, italic=False,
             color=DARK_GREY, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_bullet(doc, text, color=DARK_GREY, size=11, indent=True):
    style = 'List Bullet' if indent else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_code(doc, text):
    """Monospaced code-like paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = DARK_GREY


def make_table(doc, headers, rows, header_bg='0D47A1', col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
    for r_idx, row in enumerate(rows):
        tr  = table.rows[r_idx + 1]
        bg  = 'E3F2FD' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            p    = cell.paragraphs[0]
            run  = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GREY
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)
    return table


def ss_placeholder(doc, label: str, description: str = ''):
    """Inserts a Jira-ready screenshot placeholder box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(
        f'📷  [ SCREENSHOT: {label} ]'
        + (f'\n   → {description}' if description else '')
    )
    run.bold           = True
    run.font.size      = Pt(10)
    run.font.color.rgb = PURPLE

    # light purple border hint via styled table (1×1)
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, 'F3E5F5')
    cp = cell.paragraphs[0]
    cr = cp.add_run(
        f'JIRA ATTACHMENT PLACEHOLDER\n'
        f'Label      :  {label}\n'
        f'Description:  {description if description else "— paste screenshot here —"}\n'
        f'How to use :  Attach this image to the corresponding Jira task as test evidence.'
    )
    cr.font.name  = 'Courier New'
    cr.font.size  = Pt(9)
    cr.font.color.rgb = PURPLE
    doc.add_paragraph()


def task_card(doc, task_id: str, title: str, estimate: str, status: str,
              story: str, how: list, acceptance: list, ss_labels: list):
    """Renders one Jira-style task card with developer discussion."""
    add_heading(doc, f'{task_id}  {title}', 3, color=TEAL, space_before=12)

    meta_rows = [('Task ID', task_id), ('Title', title),
                 ('Estimate', estimate), ('Status', status)]
    t = doc.add_table(rows=len(meta_rows), cols=2)
    t.style = 'Table Grid'
    for r_idx, (k, v) in enumerate(meta_rows):
        bg = 'F0F4FF' if r_idx % 2 == 0 else 'FFFFFF'
        set_cell_bg(t.rows[r_idx].cells[0], 'E8EAF6')
        set_cell_bg(t.rows[r_idx].cells[1], bg)
        rk = t.rows[r_idx].cells[0].paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        rk.font.color.rgb = PRIMARY
        rv = t.rows[r_idx].cells[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(10)
        rv.font.color.rgb = DARK_GREY
    for row in t.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(5.3)
    doc.add_paragraph()

    # User story
    p = doc.add_paragraph()
    r = p.add_run('User Story:  ')
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = PRIMARY
    rs = p.add_run(story)
    rs.italic = True; rs.font.size = Pt(10); rs.font.color.rgb = DARK_GREY
    doc.add_paragraph()

    # Developer discussion
    add_para(doc, '🛠  Developer Implementation Notes:', bold=True,
             color=PRIMARY, size=10, space_after=3)
    for bullet in how:
        add_bullet(doc, bullet, size=10)
    doc.add_paragraph()

    # Acceptance criteria
    add_para(doc, '✅  Acceptance Criteria:', bold=True,
             color=GREEN, size=10, space_after=3)
    for ac in acceptance:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ac)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
    doc.add_paragraph()

    # Screenshot placeholders
    if ss_labels:
        add_para(doc, '📷  Test Evidence (attach to Jira task):',
                 bold=True, color=PURPLE, size=10, space_after=3)
        for ss_label, ss_desc in ss_labels:
            ss_placeholder(doc, ss_label, ss_desc)

    # divider
    p = doc.add_paragraph()
    run = p.add_run('─' * 90)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build_document() -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    title_p = doc.add_heading('PortFolioPH', 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = PRIMARY
        run.font.size = Pt(40)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run('Sprint 1 & Sprint 2 — Combined Implementation Report')
    sr.font.size = Pt(18); sr.bold = True; sr.font.color.rgb = ACCENT

    doc.add_paragraph()
    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr2 = tagline_p.add_run('Build your portfolio, own your future.')
    tr2.italic = True; tr2.font.size = Pt(13); tr2.font.color.rgb = MID_GREY

    doc.add_paragraph()
    meta_items = [
        ('Developer',     'Mark Leannie Gacutno'),
        ('Project',       'PortFolioPH – Offline-first Portfolio Builder'),
        ('Sprints',       'Sprint 1 (Week 1)  +  Sprint 2 (Week 2)'),
        ('Date',          'March 5, 2026'),
        ('Total Points',  '64 story points (32 + 32)'),
        ('Repository',    'https://github.com/auzcee/PortFolioPHH'),
        ('Status',        '✅ BOTH SPRINTS COMPLETE'),
    ]
    for label, value in meta_items:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mlr = mp.add_run(f'{label}:   ')
        mlr.bold = True; mlr.font.size = Pt(12); mlr.font.color.rgb = PRIMARY
        mvr = mp.add_run(value)
        mvr.font.size = Pt(12); mvr.font.color.rgb = DARK_GREY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS  (manual)
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        ('1.',  'Project Overview'),
        ('2.',  'Sprint 1 – Core Setup & Architecture'),
        ('2.1', 'Sprint 1 Summary'),
        ('2.2', 'Sprint 1 Task Breakdown (Developer Discussion)'),
        ('2.3', 'Sprint 1 Architecture & Schema'),
        ('2.4', 'Sprint 1 Definition of Done'),
        ('2.5', 'Sprint 1 Test Evidence (Screenshots)'),
        ('3.',  'Sprint 2 – Authentication & User Setup'),
        ('3.1', 'Sprint 2 Summary'),
        ('3.2', 'Sprint 2 Task Breakdown (Developer Discussion)'),
        ('3.3', 'Sprint 2 Bug Fixes'),
        ('3.4', 'Sprint 2 Definition of Done'),
        ('3.5', 'Sprint 2 Test Evidence (Screenshots)'),
        ('4.',  'Combined Tech Stack'),
        ('5.',  'Sprint Roadmap'),
        ('6.',  'Risks & Mitigations'),
        ('7.',  'Sign-Off'),
    ]
    for num, text in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_num = p.add_run(f'{num:<6}')
        r_num.bold = True; r_num.font.size = Pt(11); r_num.font.color.rgb = PRIMARY
        r_txt = p.add_run(text)
        r_txt.font.size = Pt(11); r_txt.font.color.rgb = DARK_GREY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '1. Project Overview', 1)
    add_para(doc, (
        'PortFolioPH is an offline-first Flutter application designed for Filipino students '
        'and fresh graduates to create, manage, and showcase professional digital portfolios '
        'directly from their devices — no internet connection required for core features.'
    ))

    make_table(doc,
        headers=['Property', 'Value'],
        rows=[
            ('Application Name',   'PortFolioPH'),
            ('Platform',           'Android (API 26–34) + Web (Flutter web dev mode)'),
            ('Architecture',       'Clean Architecture + Provider (ChangeNotifier)'),
            ('State Management',   'Provider only (ChangeNotifier)'),
            ('Database',           'SQLite via sqflite (native) / sqflite_common_ffi_web (web)'),
            ('Routing',            'GoRouter 14+'),
            ('Theme',              'Material 3, light + dark, primary #0D47A1'),
            ('Session',            'SharedPreferences (userId integer key)'),
            ('Password Security',  'SHA-256 hashing via crypto package'),
            ('Sprints Covered',    'Sprint 1 (Foundation) + Sprint 2 (Authentication)'),
        ],
        col_widths=[2.0, 4.5],
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. SPRINT 1
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '2. Sprint 1 – Core Setup & Architecture', 1, color=PRIMARY)
    doc.add_paragraph()

    # 2.1 Summary
    add_heading(doc, '2.1  Sprint 1 Summary', 2)
    add_para(doc, (
        'Sprint 1 establishes the entire technical foundation of PortFolioPH. No visible '
        'product features are shipped to end-users in this sprint — the goal is a solid, '
        'production-grade scaffold that every future sprint builds on. By the end of Sprint 1 '
        'the app launches, opens the SQLite database, checks for a saved session, and navigates '
        'to the appropriate screen.'
    ))
    make_table(doc,
        headers=['Property', 'Value'],
        rows=[
            ('Epic',          'EPIC-001 – Core Setup & Architecture'),
            ('Sprint',        'Sprint 1 (Week 1)'),
            ('Story Points',  '32 hours'),
            ('Status',        '✅ COMPLETE'),
            ('Files Created', '25+ source files'),
            ('DB Tables',     '10 (SQLite, FK enforced)'),
            ('Key Output',    'App launches, DB opens, Splash → Login routing works'),
        ],
        col_widths=[2.0, 4.5],
    )
    doc.add_paragraph()

    make_table(doc,
        headers=['Story ID', 'Title', 'Status', 'Effort'],
        rows=[
            ('STORY-001', 'Flutter Project Initialisation & Package Setup',   '✅ Done', '2 hrs'),
            ('STORY-002', 'Clean Architecture Folder Structure',               '✅ Done', '3 hrs'),
            ('STORY-003', 'SQLite DatabaseService & Schema (10 tables)',       '✅ Done', '5 hrs'),
            ('STORY-004', 'AppConstants & AppTheme (Material 3)',              '✅ Done', '3 hrs'),
            ('STORY-005', 'GoRouter Setup with All Named Routes',              '✅ Done', '4 hrs'),
            ('STORY-006', 'Bottom Navigation Scaffold + Placeholder Tabs (5)','✅ Done', '4 hrs'),
            ('STORY-007', 'Splash Screen with Session Check',                  '✅ Done', '4 hrs'),
            ('STORY-008', 'GitHub Setup, Android Permissions & Integration',   '✅ Done', '4 hrs'),
            ('STORY-009', 'Sprint 1 Documentation & Handover',                 '✅ Done', '3 hrs'),
        ],
        col_widths=[1.2, 3.6, 1.0, 0.8],
    )
    doc.add_paragraph()

    # 2.2 Task Breakdown
    add_heading(doc, '2.2  Sprint 1 Task Breakdown — Developer Discussion', 2)
    doc.add_paragraph()

    task_card(doc,
        task_id='STORY-001', title='Flutter Project Initialisation & Package Setup',
        estimate='2 hrs', status='✅ Done',
        story='AS A developer I WANT a clean Flutter project with all required dependencies SO THAT every team member starts from an identical baseline.',
        how=[
            'Ran `flutter create portfolioph` targeting Android API 26–34 (minSdk=26, targetSdk=34) and enabled web + desktop runners for flexibility.',
            'Edited pubspec.yaml to add 13 production dependencies: provider, go_router, sqflite, path_provider, path, shared_preferences, permission_handler, crypto, image_picker, cached_network_image, intl, uuid, flutter_svg.',
            'Declared asset directories assets/images/ and assets/icons/ in flutter/assets section.',
            'Ran `flutter pub get` and verified 0 solver conflicts. Committed pubspec.lock so dependency versions are reproducible.',
            'Configured analysis_options.yaml with flutter_lints and added project-specific lint rules (prefer_final_fields, avoid_print, etc.).',
        ],
        acceptance=[
            '`flutter pub get` exits with code 0, no solver conflicts.',
            'All 13 production packages appear in pubspec.lock.',
            'assets/images/ and assets/icons/ directories exist.',
            '`flutter analyze` reports 0 errors on a fresh checkout.',
        ],
        ss_labels=[
            ('STORY-001: pubspec.yaml dependencies', 'Screenshot of pubspec.yaml showing all dependencies listed'),
            ('STORY-001: flutter pub get output', 'Terminal output showing successful pub get with package list'),
        ],
    )

    task_card(doc,
        task_id='STORY-002', title='Clean Architecture Folder Structure',
        estimate='3 hrs', status='✅ Done',
        story='AS A developer I WANT a strictly enforced folder structure SO THAT presentation, domain, and data concerns never bleed into each other.',
        how=[
            'Designed a three-layer folder hierarchy under lib/: core/ (constants, exceptions, router, theme, utils), data/ (datasources, models, repositories, services), and presentation/ (providers, screens, widgets).',
            'Created all placeholder .dart files with header comments explaining each layer\'s responsibility, so the structure is self-documenting.',
            'Enforced the rule: presentation/ imports only from presentation/providers and core/; data/ never imports presentation/; core/ imports nothing from data/ or presentation/.',
            'Screens are grouped by feature (auth/, dashboard/, portfolio/) rather than by type (all screens in one flat folder) to support scaling.',
        ],
        acceptance=[
            'lib/ contains exactly three top-level subdirectories: core/, data/, presentation/.',
            'No cross-layer import violations (presentation → data skipping repositories).',
            'Every directory has at least one .dart file with a purpose-explaining header comment.',
        ],
        ss_labels=[
            ('STORY-002: Folder structure in VS Code Explorer', 'VS Code file explorer side panel showing the full lib/ tree'),
        ],
    )

    task_card(doc,
        task_id='STORY-003', title='SQLite DatabaseService & Schema (10 tables)',
        estimate='5 hrs', status='✅ Done',
        story='AS A developer I WANT a fully normalised SQLite schema created once on app install SO THAT all repositories have stable, typed tables to read/write.',
        how=[
            'Implemented DatabaseService as a Dart singleton (factory constructor + static _instance). A single shared instance guarantees only one DB connection exists at runtime.',
            'getDatabase() uses lazy initialisation (_database ??= await _open()) so the DB is only opened on first use, not at import time.',
            'Registered `PRAGMA foreign_keys = ON` in _onConfigure() — this runs on every connection open, ensuring FK constraints are enforced across the app lifecycle.',
            '_onCreate() calls _runMigration1() which wraps all 10 CREATE TABLE statements in a single db.batch() call. Batch is committed atomically — either all tables exist or none do.',
            'Defined an _onUpgrade() switch/case framework ready for v2+ migrations. Future sprints add `case 2: await _runMigration2(db); break;` without touching existing code.',
            'Created 10 tables: users, portfolios, projects, skills, education, work_experience, certifications, contacts, theme_settings, app_settings.',
            'All timestamps stored as ISO-8601 TEXT columns for portability. Booleans stored as INTEGER 0/1 (SQLite has no BOOLEAN type).',
            'Added 7 performance indexes on the most-queried foreign keys (e.g., idx_portfolios_user_id, idx_projects_portfolio_id).',
        ],
        acceptance=[
            'App cold-start creates portfolioph.db with exactly 10 tables.',
            'PRAGMA foreign_keys returns 1 on every connection.',
            'Dropping one table and rerunning onCreate re-creates all 10 atomically.',
            '`flutter analyze` reports 0 warnings on database_service.dart.',
        ],
        ss_labels=[
            ('STORY-003: DB Schema diagram', 'ERD or table list showing all 10 tables and their columns'),
            ('STORY-003: DatabaseService code structure', 'Screenshot of database_service.dart open in editor showing _runMigration1'),
        ],
    )

    task_card(doc,
        task_id='STORY-004', title='AppConstants & AppTheme (Material 3)',
        estimate='3 hrs', status='✅ Done',
        story='AS A developer I WANT all design tokens in one place SO THAT zero magic numbers or hardcoded colours appear anywhere else in the codebase.',
        how=[
            'Created AppConstants as an `abstract final class` (cannot be instantiated). Every constant is a `static const` — compile-time resolved, zero runtime cost.',
            'Grouped constants into logical sections: db (dbName, dbVersion), session (prefUserId, prefThemeMode), brand (primaryColor, accentColor, errorColor), typography (fontSizeXs→fontSizeDisplay), spacing (spacingXs=4 → spacingXxl=48), radius, elevation, animation, navigation, validation.',
            'AppTheme provides two static getters: AppTheme.light and AppTheme.dark. Both use `useMaterial3: true` with ColorScheme.fromSeed().',
            'A shared private `_buildTextTheme()` method is called by both themes to ensure identical type scale — no duplication.',
            'Styled all component themes centrally: AppBarTheme, BottomNavigationBarTheme, ElevatedButtonTheme, TextButtonTheme, InputDecorationTheme (with uniform border radius), CardTheme, DividerTheme.',
        ],
        acceptance=[
            'grep for any hex color literal (e.g. 0xFF or Color(0x)) outside app_constants.dart returns 0 results.',
            'Light and dark themes render correctly on first launch and after toggle.',
            'All spacing/radius values reference AppConstants — no raw integer literals in widget files.',
        ],
        ss_labels=[
            ('STORY-004: Light theme render', 'Screenshot of app running in Light mode showing brand colours'),
            ('STORY-004: Dark theme render', 'Screenshot of app running in Dark mode'),
        ],
    )

    task_card(doc,
        task_id='STORY-005', title='GoRouter Setup with All Named Routes',
        estimate='4 hrs', status='✅ Done',
        story='AS A developer I WANT all routes declared centrally with an auth guard SO THAT navigation is predictable and unauthenticated users cannot reach protected screens.',
        how=[
            'Created AppRoutes abstract final class enumerating all route path constants: splash (/), login (/login), register (/register), dashboard (/dashboard), plus placeholder routes for Sprint 3–6.',
            'AppRouter.create(AuthProvider authProvider) is a factory method (not a widget) that builds and returns a GoRouter instance. Keeping the router creation outside the widget tree prevents accidental recreation on rebuilds.',
            'The redirect callback checks authProvider.isAuthenticated and the current location. Unauthenticated users on protected routes → /login. Authenticated users on auth-only routes (/login, /register) → /dashboard. The splash route (/) bypasses the guard entirely.',
            'debugLogDiagnostics: true left enabled during development so route transitions are visible in the console.',
            'GoRouter is created once in _RouterScopeState.initState() via AppRouter.create(context.read<AuthProvider>()) to avoid recreation on every build.',
        ],
        acceptance=[
            'Navigating to /dashboard while logged out redirects to /login.',
            'Navigating to /login while logged in redirects to /dashboard.',
            '/ (splash) always loads regardless of auth state.',
            'All route name constants in AppRoutes are used — no hardcoded strings in navigation calls.',
        ],
        ss_labels=[
            ('STORY-005: Router redirect flow', 'Console log showing GoRouter debug output on app launch and login'),
        ],
    )

    task_card(doc,
        task_id='STORY-006', title='Bottom Navigation Scaffold + Placeholder Tabs',
        estimate='4 hrs', status='✅ Done',
        story='AS A user I WANT a 5-tab bottom navigation bar SO THAT I can switch between Home, Portfolio, Resume, Skills, and Profile without losing my scroll position.',
        how=[
            'MainScaffold wraps an IndexedStack of 5 children. IndexedStack keeps all tab widgets alive (not disposed on switch), preserving scroll position, loaded data, and text field state.',
            'NavigationProvider (ChangeNotifier) stores _currentIndex. MainScaffold reads it with context.watch<NavigationProvider>() — only the BottomNavigationBar and IndexedStack rebuild, not the whole scaffold.',
            'Each of the 5 placeholder tab bodies shows a centred icon + label + "Coming in Sprint N" message using the shared PlaceholderTabBody widget — avoiding copy-paste across 5 screens.',
            'Tab icons use Icons.home_outlined, Icons.work_outline, Icons.article_outlined, Icons.stars_outlined, Icons.person_outline to signal purpose clearly.',
        ],
        acceptance=[
            'Tapping each tab switches the visible body without losing state in other tabs.',
            'Active tab icon/label highlights in primary colour.',
            'PlaceholderTabBody renders for all 5 tabs with correct names.',
        ],
        ss_labels=[
            ('STORY-006: Bottom nav bar rendering', 'Screenshot of the main scaffold with all 5 bottom navigation tabs visible'),
            ('STORY-006: Tab state preserved', 'Two screenshots showing scroll position preserved when switching tabs'),
        ],
    )

    task_card(doc,
        task_id='STORY-007', title='Splash Screen with Session Check',
        estimate='4 hrs', status='✅ Done',
        story='AS A returning user I WANT the app to auto-login from a saved session SO THAT I do not need to enter credentials on every launch.',
        how=[
            'SplashScreen uses a single AnimationController (600ms) to fade in the logo and tagline using FadeTransition — lightweight, no heavy animation libraries needed.',
            'WidgetsBinding.addPostFrameCallback schedules _init() to run after the first frame is rendered. This prevents "setState during build" errors by ensuring the widget tree exists before navigation is triggered.',
            '_init() uses Future.wait([DatabaseService().open(), Future.delayed(splashDuration)]) — both the DB open and the minimum 3-second display run in parallel. The user always sees the splash for at least 3 seconds, but DB time does not add to it.',
            'After Future.wait completes, calls context.read<AuthProvider>().restoreSession(). This reads the stored userId from SharedPreferences and loads the matching UserModel from SQLite. If found, navigates to /dashboard; if not, navigates to /login.',
            'Error handling: the entire _init() is wrapped in try/catch. If anything fails (DB corrupt, preferences unreadable), the app gracefully routes to /login rather than crashing.',
        ],
        acceptance=[
            'App shows splash logo for ≥3 seconds on every launch.',
            'On fresh install: /login is shown after splash.',
            'On returning user: /dashboard is shown after splash without entering credentials.',
            'DB open error → splash still transitions to /login (no white screen crash).',
        ],
        ss_labels=[
            ('STORY-007: Splash screen', 'Screenshot of the splash screen showing the logo and tagline animated in'),
            ('STORY-007: Auto-login flow', 'Screenshot sequence: splash → dashboard (returning user session restored)'),
        ],
    )

    task_card(doc,
        task_id='STORY-008', title='Android Permissions & Integration Test',
        estimate='4 hrs', status='✅ Done',
        story='AS A developer I WANT all required Android permissions declared and a passing integration test SO THAT the app can be submitted to QA.',
        how=[
            'Added CAMERA, READ_EXTERNAL_STORAGE (maxSdkVersion=32), WRITE_EXTERNAL_STORAGE (maxSdkVersion=29), READ_MEDIA_IMAGES (minSdkVersion=33), and INTERNET to AndroidManifest.xml. The maxSdkVersion/minSdkVersion attributes ensure each permission is only requested on the correct Android API range.',
            'Ran `flutter analyze` and resolved all linting warnings before creating the first full integration run (Splash → Login visible on emulator).',
            'Verified app builds for Android debug APK with `flutter build apk --debug` — 0 Gradle errors.',
        ],
        acceptance=[
            'AndroidManifest.xml contains all 5 permission declarations with correct API guards.',
            '`flutter build apk --debug` exits with code 0.',
            '`flutter analyze` reports 0 errors and 0 warnings.',
        ],
        ss_labels=[
            ('STORY-008: flutter analyze output', 'Terminal showing `flutter analyze` with 0 issues'),
            ('STORY-008: Splash screen on emulator', 'Android emulator screenshot showing the splash screen running'),
        ],
    )

    doc.add_page_break()

    # 2.3 Architecture
    add_heading(doc, '2.3  Sprint 1 Architecture & Database Schema', 2)

    add_para(doc, 'Layer diagram after Sprint 1:', bold=True, color=PRIMARY, size=10)
    add_code(doc,
        '┌────────────────────────────────────────────────────────┐\n'
        '│               PRESENTATION LAYER                        │\n'
        '│  Screens (Splash, Login shell, Register shell)          │\n'
        '│  Providers: ThemeProvider, NavigationProvider,          │\n'
        '│             UserProvider (Sprint 1), PortfolioProvider  │\n'
        '│  GoRouter  ◄── AppRouter.create(UserProvider)           │\n'
        '└────────────────────────┬───────────────────────────────┘\n'
        '                         │\n'
        '┌────────────────────────▼───────────────────────────────┐\n'
        '│                  DATA LAYER                             │\n'
        '│  8 Repositories  ───►  DatabaseService (Singleton)     │\n'
        '│                               │                         │\n'
        '│                       SQLite DB (10 tables, v1)         │\n'
        '└────────────────────────────────────────────────────────┘\n'
        '  Session: SharedPreferences (userId, themeMode)'
    )
    doc.add_paragraph()

    add_heading(doc, 'Database Schema', 3, color=TEAL)
    make_table(doc,
        headers=['#', 'Table', 'PK', 'FK', 'Key Columns'],
        rows=[
            ('1',  'users',           'id AI', '—',                  'username UNIQUE, email UNIQUE, password_hash'),
            ('2',  'portfolios',      'id AI', 'users CASCADE',       'is_public, template_id, custom_url'),
            ('3',  'projects',        'id AI', 'portfolios + users',  'tech_stack, is_featured, sort_order'),
            ('4',  'skills',          'id AI', 'users CASCADE',       'category, level, years_of_experience'),
            ('5',  'education',       'id AI', 'users CASCADE',       'institution, degree, field_of_study, is_current'),
            ('6',  'work_experience', 'id AI', 'users CASCADE',       'company, job_title, employment_type, is_current'),
            ('7',  'certifications',  'id AI', 'users CASCADE',       'issuing_organization, credential_id, does_expire'),
            ('8',  'contacts',        'id AI', 'users CASCADE',       'platform, url, display_label, sort_order'),
            ('9',  'theme_settings',  'id AI', 'users CASCADE',       'theme_mode TEXT, primary_color_hex'),
            ('10', 'app_settings',    'id AI', 'users CASCADE',       'setting_key, setting_value UNIQUE(user_id,key)'),
        ],
        col_widths=[0.3, 1.3, 0.8, 1.3, 2.8],
    )
    doc.add_paragraph()

    # 2.4 DoD
    add_heading(doc, '2.4  Sprint 1 Definition of Done', 2)
    s1_dod = [
        ('✅', 'pubspec.yaml — all 13 deps present; flutter pub get = 0 errors.'),
        ('✅', 'lib/ follows Clean Architecture scaffold exactly.'),
        ('✅', 'DatabaseService singleton opens SQLite, creates 10 tables atomically.'),
        ('✅', 'PRAGMA foreign_keys = ON enforced on every connection.'),
        ('✅', 'AppConstants — zero magic numbers in any other file.'),
        ('✅', 'AppTheme — Material 3, light + dark, brand colours.'),
        ('✅', 'AppRouter — GoRouter with auth guard, all named routes.'),
        ('✅', 'main.dart — MultiProvider + MaterialApp.router wired correctly.'),
        ('✅', 'Splash fade animation, DB open, session check, redirect logic.'),
        ('✅', 'MainScaffold — IndexedStack 5 tabs, NavigationProvider driven.'),
        ('✅', 'AndroidManifest.xml — 5 permissions with API guards.'),
        ('✅', 'flutter analyze — 0 errors, 0 warnings.'),
    ]
    for icon, text in s1_dod:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        ri = p.add_run(icon + '  '); ri.font.size = Pt(11); ri.font.color.rgb = GREEN
        rt = p.add_run(text); rt.font.size = Pt(11); rt.font.color.rgb = DARK_GREY
    doc.add_paragraph()

    # 2.5 Sprint 1 Screenshots
    add_heading(doc, '2.5  Sprint 1 Test Evidence', 2)
    add_para(doc, 'The following screenshots should be attached to the corresponding Jira tasks as test evidence.', italic=True, color=MID_GREY)
    doc.add_paragraph()

    sprint1_screenshots = [
        ('Sprint 1 – App Launch (Splash Screen)',
         'Emulator/browser showing the PortFolioPH splash with logo fading in and tagline visible.'),
        ('Sprint 1 – Splash → Login Redirect (No Session)',
         'Screenshot sequence: splash finishes loading → LoginScreen appears (fresh install).'),
        ('Sprint 1 – Splash → Dashboard Redirect (With Session)',
         'Screenshot sequence: splash finishes loading → DashboardScreen (user was logged in).'),
        ('Sprint 1 – Bottom Navigation (All 5 Tabs)',
         'Screenshot showing the MainScaffold with all 5 tabs in the bottom nav bar.'),
        ('Sprint 1 – Dark Mode vs Light Mode',
         'Side-by-side or two separate screenshots showing light and dark theme rendering.'),
        ('Sprint 1 – flutter analyze 0 issues',
         'Terminal window showing `flutter analyze` output with "No issues found!" or 0 errors.'),
        ('Sprint 1 – flutter build apk debug success',
         'Terminal showing successful APK build with file size and path.'),
    ]
    for label, desc in sprint1_screenshots:
        ss_placeholder(doc, label, desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. SPRINT 2
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '3. Sprint 2 – Authentication & User Setup', 1, color=PRIMARY)
    doc.add_paragraph()

    # 3.1 Summary
    add_heading(doc, '3.1  Sprint 2 Summary', 2)
    add_para(doc, (
        'Sprint 2 delivers the complete authentication layer and initial user-profile flow. '
        'Users can register a new account, set up their profile avatar/bio, and log in. '
        'Sessions persist across app restarts via SharedPreferences. Two critical bugs '
        'discovered during testing were identified and resolved within the sprint: premature '
        'form validation errors on the register screen, and SQLite unavailability when '
        'running on Flutter web.'
    ))
    make_table(doc,
        headers=['Property', 'Value'],
        rows=[
            ('Epic',          'EPIC-002 – Authentication & User Setup'),
            ('Sprint',        'Sprint 2 (Week 2)'),
            ('Story Points',  '32 hours'),
            ('Status',        '✅ COMPLETE'),
            ('New Files',     '6 (AuthException, AuthService, ProfileService, AuthProvider, ProfileSetupScreen, validators.dart, date_formatter.dart)'),
            ('Updated Files', '8 (RegisterScreen, LoginScreen, DashboardScreen, ProfileScreen, SplashScreen, AppRouter, main.dart, DatabaseService)'),
            ('Bug Fixes',     '2 (Form validation timing, Web SQLite support)'),
            ('Key Output',    'Full flow: Register → Profile Setup → Dashboard → Logout → Login'),
        ],
        col_widths=[2.0, 4.5],
    )
    doc.add_paragraph()

    make_table(doc,
        headers=['Task #', 'Title', 'Status', 'Effort'],
        rows=[
            ('S2-01', 'AuthException typed exception class',                   '✅ Done', '1 h'),
            ('S2-02', 'AuthService.register() — validate, hash, INSERT',       '✅ Done', '3 h'),
            ('S2-03', 'AuthService.login() — findByEmail + SHA-256 compare',   '✅ Done', '2 h'),
            ('S2-04', 'AuthProvider ChangeNotifier',                           '✅ Done', '3 h'),
            ('S2-05', 'RegisterScreen — full 5-field form w/ real-time validation', '✅ Done', '4 h'),
            ('S2-06', 'LoginScreen — wired to AuthProvider',                   '✅ Done', '2 h'),
            ('S2-07', 'ProfileService — getProfile, updateProfile, updateAvatar', '✅ Done', '2 h'),
            ('S2-08', 'ProfileSetupScreen — avatar, bio, school, year level',  '✅ Done', '3 h'),
            ('S2-09', 'DashboardScreen — greeting + stat cards + quick actions', '✅ Done', '3 h'),
            ('S2-10', 'Session persistence & restoreSession()',                '✅ Done', '2 h'),
            ('S2-11', 'AppValidators + AppDateFormatter utility classes',       '✅ Done', '2 h'),
            ('BUG-01', 'FIX: Register form premature validation errors',        '✅ Fixed', '1 h'),
            ('BUG-02', 'FIX: Registration fails on Flutter web (no SQLite)',    '✅ Fixed', '3 h'),
        ],
        col_widths=[0.8, 3.8, 0.8, 0.6],
    )
    doc.add_paragraph()

    # 3.2 Task Breakdown
    add_heading(doc, '3.2  Sprint 2 Task Breakdown — Developer Discussion', 2)
    doc.add_paragraph()

    task_card(doc,
        task_id='S2-01', title='AuthException Typed Exception Class',
        estimate='1 hr', status='✅ Done',
        story='AS A developer I WANT a typed auth exception SO THAT UI layers can display specific error messages without parsing generic exception strings.',
        how=[
            'Created lib/core/exceptions/auth_exception.dart with `class AuthException implements Exception`.',
            'Two fields: `message` (required String) and `code` (optional String). The code field carries machine-readable identifiers like "email_taken", "invalid_credentials", "insert_failed" — allowing UI to branch on specific failure modes if needed.',
            'Both fields are final (immutable). The class is `const`-constructible for efficiency.',
            'AuthService throws typed AuthException everywhere; AuthProvider catches it specifically in the first catch block. The second `catch (e)` handles unexpected exceptions — this two-tier pattern ensures AuthException messages always surface correctly.',
        ],
        acceptance=[
            'AuthException can be constructed as `const AuthException("msg", code: "code")`.',
            'AuthProvider.register() catch block displays e.message (not a generic string) when AuthException is thrown.',
            'No `catch (e)` swallows an AuthException (verified by code review).',
        ],
        ss_labels=[
            ('S2-01: AuthException class code', 'Screenshot of auth_exception.dart open in editor'),
        ],
    )

    task_card(doc,
        task_id='S2-02', title='AuthService.register() — Validate, Hash, INSERT',
        estimate='3 hrs', status='✅ Done',
        story='AS A new user I WANT to create an account with a unique username and email SO THAT my data is saved securely.',
        how=[
            'AuthService.register() takes username, email, password (all required) and fullName (optional). All parameters are received pre-validated by the form; the service also runs server-side guards.',
            'Field-level guards (blank checks) run first and throw AuthException immediately. This prevents unnecessary DB calls.',
            'Two async uniqueness checks run sequentially: findByEmail() then findByUsername(). Both query the SQLite DB using parameterised WHERE clauses — no string interpolation.',
            'Password is hashed with AppHelpers.hashPassword(password) — which calls SHA-256 from the crypto package and returns a hex string. The plain-text password is never stored or returned.',
            'A new UserModel is built with all fields normalised (email.trim().toLowerCase(), username.trim()) and createdAt/updatedAt stamped with AppHelpers.nowIso().',
            'UserRepository.insert() runs inside a try/catch. Any DB-level failure (constraint violation, disk full) is caught and rethrown as `AuthException(code: "insert_failed")`.',
        ],
        acceptance=[
            'Registering with an existing email throws AuthException(code: "email_taken").',
            'Registering with an existing username throws AuthException(code: "username_taken").',
            'The stored password_hash is a 64-character hex string, never the plaintext.',
            'Registered user\'s email is stored in lowercase.',
        ],
        ss_labels=[
            ('S2-02: AuthService.register() code', 'Screenshot of auth_service.dart register() method open in editor'),
            ('S2-02: Successful registration in DB', 'SQLite browser or debug print showing inserted user row with hashed password'),
        ],
    )

    task_card(doc,
        task_id='S2-03', title='AuthService.login() — findByEmail + SHA-256 Compare',
        estimate='2 hrs', status='✅ Done',
        story='AS A returning user I WANT to log in with my email and password SO THAT I can access my portfolio data.',
        how=[
            'login() calls UserRepository.findByEmail() which performs a case-insensitive lookup (email stored normalised). Returns null if not found.',
            'On null result, throws AuthException("Invalid email or password.", code: "invalid_credentials") — the SAME message used when the password is wrong. This prevents email enumeration attacks: an attacker cannot determine whether an email is registered by observing different error messages.',
            'On a found user, re-hashes the input password with AppHelpers.hashPassword() and compares to the stored hash using `!=` equality. If they differ, same generic AuthException is thrown.',
            'On match, returns the full UserModel. The caller (AuthProvider) then persists the session.',
        ],
        acceptance=[
            'login() with wrong email: throws AuthException("Invalid email or password.").',
            'login() with correct email but wrong password: same message (no difference to user).',
            'login() with correct credentials: returns UserModel with all fields populated.',
        ],
        ss_labels=[
            ('S2-03: Login screen working', 'Screenshot of LoginScreen with valid credentials before tapping Log In'),
            ('S2-03: Login success → Dashboard', 'Screenshot of DashboardScreen appearing after successful login'),
            ('S2-03: Login error SnackBar', 'Screenshot of SnackBar showing "Invalid email or password."'),
        ],
    )

    task_card(doc,
        task_id='S2-04', title='AuthProvider ChangeNotifier',
        estimate='3 hrs', status='✅ Done',
        story='AS A UI developer I WANT a reactive auth state provider SO THAT all widgets depending on login state rebuild automatically on auth changes.',
        how=[
            'AuthProvider extends ChangeNotifier. It is provided at the root MultiProvider in main.dart so every widget in the tree can access it.',
            '_begin() is an internal helper called at the start of every async action: sets _isLoading = true, clears _errorMessage, calls notifyListeners(). This ensures the UI immediately shows a loading indicator.',
            '_endLoading() is called in the `finally` block of every async action: sets _isLoading = false, notifyListeners(). Using `finally` guarantees this runs even if an exception is thrown.',
            'Every mutation (currentUser assignment, errorMessage assignment) is followed by notifyListeners() to trigger rebuilds.',
            '_persistSession(int userId) saves the userId integer to SharedPreferences using the key AppConstants.prefUserId.',
            'restoreSession() reads the key from SharedPreferences; if null, returns false. If found, loads UserModel from UserRepository.findById(). If the user row is gone (e.g. DB deleted), removes the stale preference and returns false.',
            'updateCurrentUser(UserModel) is called by ProfileSetupScreen after saving — updates the in-memory currentUser without re-querying the DB.',
        ],
        acceptance=[
            'context.watch<AuthProvider>().isLoading rebuilds the submit button\'s enabled state in real-time.',
            'After logout(), currentUser is null and isAuthenticated is false.',
            'After app restart with saved session, restoreSession() returns true and currentUser is populated.',
        ],
        ss_labels=[
            ('S2-04: AuthProvider state flow', 'Screenshot of auth_provider.dart register() method in editor'),
        ],
    )

    task_card(doc,
        task_id='S2-05', title='RegisterScreen — 5-Field Form with Real-Time Validation',
        estimate='4 hrs', status='✅ Done',
        story='AS A new user I WANT inline validation on the registration form SO THAT I know exactly what is wrong before I submit.',
        how=[
            'Five TextFormFields: Full Name (optional), Username, Email, Password, Confirm Password.',
            'Form uses `autovalidateMode: AutovalidateMode.onUserInteraction` — each field shows its error only after the user has typed in it. This prevents the initial state showing all fields as invalid.',
            'Five TextEditingControllers are created in _RegisterScreenState. All five controllers are registered with addListener(_onFieldChanged) in initState().',
            '_onFieldChanged() computes form validity by calling each AppValidators function directly on the raw text values. This is a pure, side-effect-free check — it does NOT call FormState.validate(), which would force all fields into error state immediately.',
            'The `_formValid` bool gate is used to enable/disable the Create Account ElevatedButton. The button is disabled while `isLoading || !_formValid`.',
            'On submit, _submit() calls AuthProvider.register(). On success → context.go("/profile-setup"). On failure → SnackBar with auth.errorMessage.',
            'Password visibility toggle (eye icon button) implemented for both Password and Confirm Password fields.',
        ],
        acceptance=[
            'Fresh page load: no fields show red error borders.',
            'Typing invalid username then moving away: only username field turns red.',
            'All fields valid: Create Account button becomes enabled (blue).',
            'Duplicate email: SnackBar shows "An account with this email already exists."',
        ],
        ss_labels=[
            ('S2-05: Register screen clean state', 'Screenshot of register screen immediately after navigating to it (no errors)'),
            ('S2-05: Register inline validation', 'Screenshot showing one or more fields with inline error messages after user interaction'),
            ('S2-05: Register button enabled', 'Screenshot with all fields valid and Create Account button actively enabled/blue'),
            ('S2-05: Register duplicate email error', 'Screenshot of SnackBar error after attempting to register with an already-used email'),
        ],
    )

    task_card(doc,
        task_id='S2-06', title='LoginScreen — Wired to AuthProvider',
        estimate='2 hrs', status='✅ Done',
        story='AS A returning user I WANT to log in with my email and password SO THAT my session is restored.',
        how=[
            'Updated the Sprint 1 LoginScreen shell to use AuthProvider instead of UserProvider. Import path changed accordingly.',
            'Validator callbacks changed from inline lambdas to AppValidators.validateEmail and AppValidators.validatePassword — consistent with RegisterScreen.',
            'The _submit() async method reads context.read<AuthProvider>() and calls auth.login(email: ..., password: ...). The `await` ensures _isLoading is true during the request, disabling the Login button.',
            'SnackBar on failure reads auth.errorMessage which comes from AuthException.message (the generic "Invalid email or password." string).',
            'On success → context.go("/dashboard").',
        ],
        acceptance=[
            'Login with correct credentials → DashboardScreen.',
            'Login with wrong password → SnackBar "Invalid email or password.".',
            'Login button shows CircularProgressIndicator while request is in flight.',
        ],
        ss_labels=[
            ('S2-06: Login screen UI', 'Screenshot of the Login screen in clean state'),
            ('S2-06: Login loading state', 'Screenshot of the Login screen with CircularProgressIndicator inside the button'),
            ('S2-06: Login error SnackBar', 'Screenshot of the error SnackBar at the bottom of the screen'),
            ('S2-06: Login success Dashboard', 'Screenshot of DashboardScreen appearing after successful login'),
        ],
    )

    task_card(doc,
        task_id='S2-07', title='ProfileService — getProfile, updateProfile, updateAvatar',
        estimate='2 hrs', status='✅ Done',
        story='AS A developer I WANT a dedicated service for profile operations SO THAT auth concerns (AuthService) and profile edit concerns are separated per SRP.',
        how=[
            'ProfileService follows the same injectable constructor pattern as AuthService: `ProfileService({UserRepository? userRepository})` — defaults to the singleton but allows test injection.',
            'getProfile(int userId) is a simple delegation to UserRepository.findById(). Returns null if the user does not exist.',
            'updateProfile(UserModel user) stamps `updatedAt = AppHelpers.nowIso()` before calling UserRepository.update(). Returns the saved model with the new timestamp set.',
            'updateAvatar(UserModel user, String avatarPath) is a convenience wrapper: calls updateProfile(user.copyWith(avatarPath: avatarPath)) — no duplication of update logic.',
        ],
        acceptance=[
            'updateProfile() saves changes to SQLite and the returned model has the new updatedAt timestamp.',
            'Calling getProfile(userId) after updateProfile() returns the updated data.',
            'AuthService imports are absent from ProfileService (services are fully separated).',
        ],
        ss_labels=[
            ('S2-07: ProfileService code', 'Screenshot of profile_service.dart open in editor'),
        ],
    )

    task_card(doc,
        task_id='S2-08', title='ProfileSetupScreen — Avatar, Bio, School, Year Level',
        estimate='3 hrs', status='✅ Done',
        story='AS A new user I WANT to set up my profile photo and academic details right after registering SO THAT my portfolio has personal context from day one.',
        how=[
            'ProfileSetupScreen is shown at /profile-setup immediately after successful registration.',
            'Avatar picker calls image_picker with ImageSource.camera or ImageSource.gallery via an AlertDialog. Image dimensions are constrained to maxWidth=512, maxHeight=512, imageQuality=85 to keep storage efficient.',
            'kIsWeb guard: on web, ImagePicker returns an XFile whose path is a blob URL — displayed with Image.network(). On native, displayed with Image.file(File(path)).',
            'Initials fallback: if no avatar is selected, a CircleAvatar shows the first letter of the user\'s name (from AuthProvider.currentUser).',
            'Bio is a TextFormField with maxLines=3, maxLength=AppConstants.maxBioLength (500 chars) with the built-in Flutter character counter.',
            'Year Level is a DropdownButtonFormField with values: 1st Year, 2nd Year, 3rd Year, 4th Year, Graduate.',
            'Save button calls ProfileService.updateProfile(user.copyWith(bio: ..., location: ...)) then AuthProvider.updateCurrentUser(saved) to update in-memory state. Skip button in the AppBar navigates directly to /dashboard.',
        ],
        acceptance=[
            'Avatar picker opens when tapping the avatar circle.',
            'Tapping Skip navigates to /dashboard (profile fields not required).',
            'Tapping Save with bio filled → /dashboard, DashboardScreen shows the entered name.',
            'Avatar image persists for the duration of the session.',
        ],
        ss_labels=[
            ('S2-08: Profile setup screen', 'Screenshot of the full ProfileSetupScreen with all fields visible'),
            ('S2-08: Avatar picker dialog', 'Screenshot of the Camera / Gallery picker AlertDialog'),
            ('S2-08: Profile with avatar selected', 'Screenshot showing a selected avatar image in the CircleAvatar'),
        ],
    )

    task_card(doc,
        task_id='S2-09', title='DashboardScreen — Greeting + Stat Cards + Quick Actions',
        estimate='3 hrs', status='✅ Done',
        story='AS A logged-in user I WANT a personalised dashboard SO THAT I see a summary of my portfolio data at a glance.',
        how=[
            'DashboardScreen reads AuthProvider.currentUser via context.watch<AuthProvider>() so it rebuilds automatically when the user profile is updated.',
            'Greeting card: "Hello, [fullName]!" (falling back to username if fullName is null). AppConstants.appTagline shown below.',
            'Stats GridView (2×2): Portfolios, Projects, Skills, Education — each showing count=0 in Sprint 2 (real counts wired from DB in Sprint 3). Built as a private _StatCard widget with icon, count (headline style), and label.',
            'Quick Actions Card: three _ActionTile rows — Create Portfolio, Export Resume, Share Portfolio. Each shows a "Coming in Sprint N" subtitle and a chevron. Tapping shows a SnackBar placeholder.',
            'Kept the two private widget classes (_StatCard, _ActionTile) at the bottom of the file to avoid creating separate files for such tightly coupled widgets.',
        ],
        acceptance=[
            'Dashboard shows the correct user\'s name from AuthProvider.currentUser.',
            'Four stat cards render with 0 counts (stubbed until Sprint 3).',
            'Three quick action tiles are visible in the card.',
            'Logout from ProfileScreen clears the name and redirects to /login.',
        ],
        ss_labels=[
            ('S2-09: Dashboard screen', 'Screenshot of the full DashboardScreen showing greeting, stat cards, and quick actions'),
            ('S2-09: Dashboard after profile setup', 'Screenshot of dashboard showing the user\'s name entered during profile setup'),
        ],
    )

    task_card(doc,
        task_id='S2-10', title='Session Persistence & restoreSession()',
        estimate='2 hrs', status='✅ Done',
        story='AS A user I WANT the app to remember my login SO THAT I do not have to re-enter credentials on every launch.',
        how=[
            'AuthProvider._persistSession(int userId) writes the integer userId to SharedPreferences using key AppConstants.prefUserId ("user_id"). This is called on both successful register and login.',
            'AuthProvider.restoreSession() is called in SplashScreen._init() after the database is opened. It reads the saved userId. If absent → returns false. If present → loads UserModel from UserRepository.findById(userId).',
            'If findById returns null (user record was deleted since the session was saved), restoreSession() removes the stale key from SharedPreferences and returns false — avoiding a state where isAuthenticated=true but currentUser=null.',
            'AuthProvider.logout() removes the prefUserId key and sets _currentUser = null. The GoRouter redirect guard detects isAuthenticated = false and redirects all protected routes to /login.',
        ],
        acceptance=[
            'Login → close app → reopen: user is taken to /dashboard directly (session restored).',
            'Logout → close app → reopen: user is taken to /login (session cleared).',
            'Corrupt SharedPreferences / missing user row: app gracefully routes to /login.',
        ],
        ss_labels=[
            ('S2-10: Session restore test', 'Two screenshots: (1) app launched with saved session → dashboard, (2) after logout → login screen on next launch'),
        ],
    )

    task_card(doc,
        task_id='S2-11', title='AppValidators & AppDateFormatter Utility Classes',
        estimate='2 hrs', status='✅ Done',
        story='AS A developer I WANT centralised validation and date formatting functions SO THAT every form and date display uses consistent, tested logic.',
        how=[
            'AppValidators is an `abstract final class` (non-instantiable) with 6 static validator functions: validateEmail, validatePassword, validateConfirmPassword, validateUsername, validateRequired, validateOptionalUrl.',
            'Every function matches the FormField.validator type signature: `String? Function(String?)` — they can be passed directly to the validator: parameter of TextFormField without wrapping.',
            'AppDateFormatter is similarly an `abstract final class` with 8 static functions: formatDate, formatShort, formatMonthYear, formatFull, formatTime, formatDateTime, formatDateRange, formatRelative.',
            'All formatter functions return "—" on null, empty, or unparseable input — no throw, no crash. This is important because historical data may have missing date fields.',
            'formatRelative() implements a tiered relative time algorithm: "just now" (<1 min), "Xm ago" (<1 h), "Xh ago" (<24 h), "Xd ago" (<7 d), "Xw ago" (<28 d), "Xmo ago" (<12 mo), "Xy ago" (≥12 mo).',
        ],
        acceptance=[
            'AppValidators.validateEmail(null) returns "Email is required."',
            'AppValidators.validateUsername("a!") returns the invalid character error.',
            'AppValidators.validateConfirmPassword("abc", "xyz") returns "Passwords do not match."',
            'AppDateFormatter.formatRelative(nowIso()) returns "just now".',
            'AppDateFormatter.formatDateRange(startIso, null) returns "Jan 2026 – Present".',
        ],
        ss_labels=[
            ('S2-11: AppValidators code', 'Screenshot of validators.dart showing the 6 validator functions'),
            ('S2-11: AppDateFormatter code', 'Screenshot of date_formatter.dart showing the formatter functions'),
        ],
    )

    doc.add_paragraph()

    # 3.3 Bug Fixes
    add_heading(doc, '3.3  Sprint 2 Bug Fixes', 2)
    doc.add_paragraph()

    task_card(doc,
        task_id='BUG-01',
        title='RegisterScreen — Premature Form Validation Errors',
        estimate='1 hr', status='✅ Fixed',
        story='AS A user I WANT the register form to only show errors on fields I have already touched SO THAT I am not confronted with red error borders on fields I have not yet typed in.',
        how=[
            'ROOT CAUSE: _onFieldChanged() was calling `_formKey.currentState?.validate()`. FormState.validate() is a "validate all registered fields now" call — it marks every field as dirty regardless of whether the user has visited it. This completely bypassed autovalidateMode: AutovalidateMode.onUserInteraction.',
            'Additionally, the Password TextFormField had `onChanged: (_) => _formKey.currentState?.validate()` — a second, redundant call that triggered the same bug on every password keystroke.',
            'FIX: Replaced the FormState.validate() call inside _onFieldChanged() with a silent computation: call each AppValidators function directly on the raw .text value and AND the results together to produce the _formValid bool. This has zero side-effects on the Form widget.',
            'FIX: Removed the `onChanged: (_) => _formKey.currentState?.validate()` line from the Password field entirely. The password controller already has _onFieldChanged registered via initState, making this call redundant.',
            'The net effect: _formValid (used to gate the submit button) is always up to date, autovalidateMode: onUserInteraction works as designed — each field shows its error only after the user has typed in it.',
        ],
        acceptance=[
            'Fresh load of /register: all fields have clean borders (no red).',
            'Type an invalid username, press Tab: only the username field shows red.',
            'Fill all fields correctly: Create Account button becomes active without triggering errors on untouched fields.',
            'Reopening /register after a failed attempt: form resets to clean state.',
        ],
        ss_labels=[
            ('BUG-01: BEFORE — premature errors', 'Screenshot from before the fix: register screen with all fields showing red error text immediately on load'),
            ('BUG-01: AFTER — clean initial state', 'Screenshot after the fix: register screen freshly opened, no fields in error state'),
            ('BUG-01: AFTER — per-field validation', 'Screenshot after the fix: only the field the user has typed in shows its error'),
        ],
    )

    task_card(doc,
        task_id='BUG-02',
        title='Registration Fails on Flutter Web — "Registration failed. Please try again."',
        estimate='3 hrs', status='✅ Fixed',
        story='AS A developer testing on Flutter web I WANT the app to work on localhost SO THAT I can demo and test features without an Android emulator.',
        how=[
            'ROOT CAUSE: DatabaseService.getDatabase() had `if (kIsWeb) throw UnsupportedError(...)`. The web target was running at localhost:51994. When AuthProvider.register() called AuthService.register() → UserRepository.insert() → DatabaseService.getDatabase(), it threw UnsupportedError. This was NOT an AuthException, so it fell into the generic `catch(e)` in AuthProvider, surfacing as the generic "Registration failed. Please try again." message — hiding the actual error.',
            'FIX STEP 1: Added `sqflite_common_ffi_web: ^1.1.1` to pubspec.yaml dependencies.',
            'FIX STEP 2: Ran `dart run sqflite_common_ffi_web:setup` to generate sqlite3.wasm and sqflite_sw.js in the web/ directory. These files are the IndexedDB-backed SQLite engine the browser uses.',
            'FIX STEP 3: In main.dart, imported `package:sqflite/sqflite.dart` (exposes the databaseFactory setter) and `package:sqflite_common_ffi_web/sqflite_ffi_web.dart` (exposes databaseFactoryFfiWeb). Added `if (kIsWeb) databaseFactory = databaseFactoryFfiWeb;` before runApp().',
            'FIX STEP 4: In DatabaseService, removed the UnsupportedError throw from getDatabase(). Updated _open() to use a bare filename (the IndexedDB key) on web vs. a full filesystem path on native, using a `kIsWeb` branch.',
            'COMPILE ERROR MID-FIX: `databaseFactory` setter was unknown because only the ffi_web import was added without sqflite.dart. Fixed by adding `import package:sqflite/sqflite.dart` to main.dart.',
            'OUTCOME: SQLite data is stored in the browser\'s IndexedDB on web — persisted across page reloads, cross-tab safe (shared worker handles locking), with identical query API to the native sqflite driver.',
        ],
        acceptance=[
            'Running `flutter run -d chrome` compiles without errors.',
            'Navigating to /register, filling all fields, and tapping Create Account creates a user row.',
            'Navigating to /login with the same credentials and tapping Log In shows the Dashboard.',
            'Refreshing the browser tab: session is restored (kIsWeb path in restoreSession() works).',
            'On Android: behaviour unchanged — still uses native sqflite with filesystem path.',
        ],
        ss_labels=[
            ('BUG-02: BEFORE — Registration failed error', 'Screenshot of the register screen showing the "Registration failed. Please try again." SnackBar (before fix)'),
            ('BUG-02: AFTER — Successful registration', 'Screenshot of the register screen after fixing: all fields valid, Create Account tapped, and ProfileSetupScreen appearing'),
            ('BUG-02: AFTER — Login success on web', 'Screenshot of DashboardScreen on Flutter web (localhost) after logging in with registered credentials'),
            ('BUG-02: web/sqlite3.wasm + sqflite_sw.js files', 'Screenshot of the web/ folder in VS Code Explorer showing sqlite3.wasm and sqflite_sw.js present'),
        ],
    )

    doc.add_page_break()

    # 3.4 DoD Sprint 2
    add_heading(doc, '3.4  Sprint 2 Definition of Done', 2)
    s2_dod = [
        ('✅', 'AuthException — typed exception, const-constructible, code field present.'),
        ('✅', 'AuthService.register() — validates, checks uniqueness, hashes, inserts.'),
        ('✅', 'AuthService.login() — generic error (no email enumeration), SHA-256 compare.'),
        ('✅', 'ProfileService — getProfile, updateProfile, updateAvatar — separate from AuthService.'),
        ('✅', 'AuthProvider — all mutations call notifyListeners(); _begin()/_endLoading() in finally.'),
        ('✅', 'RegisterScreen — 5 fields, no premature errors, button gated, SnackBar on failure.'),
        ('✅', 'LoginScreen — AuthProvider.login(), loading state, SnackBar on error.'),
        ('✅', 'ProfileSetupScreen — image_picker, bio, year level, Skip supported.'),
        ('✅', 'DashboardScreen — personalised greeting, 4-stat cards, 3-action tiles.'),
        ('✅', 'Session persists across app restarts (SharedPreferences userId).'),
        ('✅', 'AppValidators — 6 functions matching FormField.validator signature.'),
        ('✅', 'AppDateFormatter — 8 functions, "—" fallback on null/error.'),
        ('✅', 'AppRouter updated to AuthProvider; /profile-setup route added.'),
        ('✅', 'main.dart — AuthProvider in MultiProvider; kIsWeb → databaseFactoryFfiWeb.'),
        ('✅', 'DatabaseService — UnsupportedError removed; kIsWeb filename branch.'),
        ('✅', 'web/sqlite3.wasm + web/sqflite_sw.js present for web builds.'),
        ('✅', 'BUG-01 fixed — no premature form errors on RegisterScreen.'),
        ('✅', 'BUG-02 fixed — registration works on Flutter web.'),
        ('✅', 'flutter analyze — 0 errors, 0 warnings.'),
        ('✅', 'Full flow verified: Register → Profile Setup → Dashboard → Logout → Login.'),
    ]
    for icon, text in s2_dod:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        ri = p.add_run(icon + '  '); ri.font.size = Pt(11); ri.font.color.rgb = GREEN
        rt = p.add_run(text); rt.font.size = Pt(11); rt.font.color.rgb = DARK_GREY
    doc.add_paragraph()

    # 3.5 Sprint 2 Screenshots
    add_heading(doc, '3.5  Sprint 2 Test Evidence (Screenshots)', 2)
    add_para(doc, 'Attach each screenshot to the corresponding Jira task. Use the label as the attachment file name for traceability.', italic=True, color=MID_GREY)
    doc.add_paragraph()

    add_heading(doc, 'Auth Flow Screenshots', 3, color=TEAL)
    sprint2_auth_ss = [
        ('Sprint2-Register-CleanState',
         'Register screen immediately after navigating to /register — all fields clean, no red borders, Create Account button disabled (greyed).'),
        ('Sprint2-Register-InlineValidation',
         'Register screen after user has typed an invalid value in a field and moved focus — only that field shows its error. All other untouched fields remain clean.'),
        ('Sprint2-Register-AllValid',
         'Register screen with all 5 fields filled in correctly — Create Account button is blue and enabled.'),
        ('Sprint2-Register-DuplicateEmail-Error',
         'SnackBar at the bottom: "An account with this email already exists." (email already registered). Create Account button is not loading.'),
        ('Sprint2-Register-Success-ProfileSetup',
         'ProfileSetupScreen appearing after a successful registration (navigated from /register to /profile-setup).'),
        ('Sprint2-ProfileSetup-Screen',
         'Full ProfileSetupScreen: CircleAvatar with initials, Bio text field, School, and Year Level dropdown visible.'),
        ('Sprint2-ProfileSetup-AvatarPicker',
         'AlertDialog showing Camera and Gallery options when tapping the avatar button.'),
        ('Sprint2-Dashboard-Greeting',
         'DashboardScreen showing the personalised greeting with the registered user\'s name.'),
        ('Sprint2-Dashboard-StatCards',
         'DashboardScreen stat grid showing 4 cards with 0 counts (Portfolios, Projects, Skills, Education).'),
        ('Sprint2-Login-Screen',
         'LoginScreen in clean state with Email and Password fields empty.'),
        ('Sprint2-Login-Success',
         'DashboardScreen after successful login from a previously registered account.'),
        ('Sprint2-Login-InvalidCredentials',
         'SnackBar showing "Invalid email or password." after entering wrong password.'),
        ('Sprint2-Logout-BackToLogin',
         'ProfileScreen showing the Logout button, then LoginScreen after tapping Logout.'),
    ]
    for label, desc in sprint2_auth_ss:
        ss_placeholder(doc, label, desc)

    doc.add_paragraph()
    add_heading(doc, 'Bug Fix Screenshots', 3, color=RED)
    bug_ss = [
        ('BUG-01-BEFORE-PrematureErrors',
         'BEFORE FIX: Register screen showing red error borders on Email, Password, and Confirm Password fields immediately on page load — before the user has typed anything.'),
        ('BUG-01-AFTER-CleanState',
         'AFTER FIX: Register screen showing clean state on load. Zero red borders. Only the Username field which the user typed in shows its validation error.'),
        ('BUG-02-BEFORE-RegistrationFailed',
         'BEFORE FIX: Register screen showing "Registration failed. Please try again." SnackBar after tapping Create Account on Flutter web (localhost).'),
        ('BUG-02-AFTER-WebRegistrationSuccess',
         'AFTER FIX: ProfileSetupScreen appearing after successful registration on Flutter web (localhost:PORT) — SQLite via IndexedDB working.'),
        ('BUG-02-WebFolder-WasmFiles',
         'web/ folder in VS Code Explorer showing sqlite3.wasm and sqflite_sw.js files present alongside index.html.'),
    ]
    for label, desc in bug_ss:
        ss_placeholder(doc, label, desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. COMBINED TECH STACK
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '4. Combined Tech Stack (Sprint 1 + Sprint 2)', 1)

    make_table(doc,
        headers=['Category', 'Package', 'Version', 'Sprint', 'Usage'],
        rows=[
            ('State',       'provider',                  '^6.1.2',    'S1', 'ChangeNotifier state across app'),
            ('Routing',     'go_router',                 '^14.3.0',   'S1', 'Named routes, auth guard'),
            ('DB (native)', 'sqflite',                   '^2.3.3+1',  'S1', 'SQLite offline storage'),
            ('DB (web)',    'sqflite_common_ffi_web',    '^1.1.1',    'S2', 'SQLite over IndexedDB for Flutter web'),
            ('DB path',     'path_provider',             '^2.1.4',    'S1', 'Documents dir for DB file'),
            ('DB path',     'path',                      '^1.9.0',    'S1', 'join() path helper'),
            ('Session',     'shared_preferences',        '^2.3.3',    'S1', 'userId + themeMode persistence'),
            ('Permissions', 'permission_handler',        '^11.3.1',   'S1', 'Camera / storage runtime requests'),
            ('Security',    'crypto',                    '^3.0.5',    'S1', 'SHA-256 password hashing'),
            ('Media',       'image_picker',              '^1.1.2',    'S2', 'Camera/gallery avatar selection'),
            ('Media',       'cached_network_image',      '^3.4.1',    'S1', 'CDN image caching (Sprint 3+)'),
            ('Utilities',   'intl',                      '^0.20.2',   'S2', 'Date formatting via DateFormat'),
            ('Utilities',   'uuid',                      '^4.5.1',    'S1', 'UUID generation'),
            ('UI',          'flutter_svg',               '^2.0.10+1', 'S1', 'SVG icon support'),
            ('UI',          'cupertino_icons',           '^1.0.8',    'S1', 'iOS-style icon set'),
        ],
        col_widths=[1.0, 1.9, 1.0, 0.5, 2.2],
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. SPRINT ROADMAP
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '5. Sprint Roadmap', 1)

    make_table(doc,
        headers=['Sprint', 'Focus Area', 'Key Deliverables', 'Status'],
        rows=[
            ('Sprint 1', 'Core Setup & Architecture',
             'DB schema (10 tables), Clean Architecture scaffold, GoRouter, Material 3 theme, Splash',
             '✅ Complete'),
            ('Sprint 2', 'Authentication & User Setup',
             'AuthService, AuthProvider, Register/Login forms, ProfileSetup, Dashboard, Web SQL fix',
             '✅ Complete'),
            ('Sprint 3', 'Portfolio & Projects CRUD',
             'Create/Edit Portfolio, Add Projects, Project Detail screen, PortfolioProvider full CRUD',
             '🔜 Next'),
            ('Sprint 4', 'Resume Builder',
             'Education, Work Experience, Certifications — add/edit/delete with date pickers',
             '🔜'),
            ('Sprint 5', 'Skills Management',
             'Skill categories, proficiency levels, sorting, category filter chips',
             '🔜'),
            ('Sprint 6', 'Profile Edit + Settings',
             'Edit all profile fields, theme toggle, privacy settings, account deletion',
             '🔜'),
            ('Sprint 7', 'Export & Sharing',
             'PDF resume generation, Share portfolio link, QR code support',
             '🔜'),
            ('Sprint 8', 'Polish, Testing & Release',
             'Unit + widget tests, CI/CD GitHub Actions, profiling, Release APK',
             '🔜'),
        ],
        col_widths=[0.8, 1.8, 3.2, 0.8],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. RISKS & MITIGATIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '6. Risks & Mitigations', 1)

    make_table(doc,
        headers=['Risk', 'Sprint', 'Impact', 'Mitigation', 'Status'],
        rows=[
            ('Flutter/Android SDK mismatch',
             'S1', 'High',
             'Locked API 26–34 in build.gradle; flutter doctor verified.',
             '✅ Resolved'),
            ('Schema changes post-Sprint 1',
             'S1', 'Med',
             'ERD finalised before coding; _onUpgrade() migration framework ready.',
             '✅ Mitigated'),
            ('GoRouter auth guard loop',
             'S1', 'Low',
             'Splash route excluded from redirect; tested fresh install + logged-in.',
             '✅ Resolved'),
            ('sqflite unsupported on web',
             'S2', 'High',
             'sqflite_common_ffi_web added; databaseFactory overridden on kIsWeb in main(); sqlite3.wasm + sqflite_sw.js in web/.',
             '✅ Resolved'),
            ('Premature form validation (UX)',
             'S2', 'Med',
             'Replaced FormState.validate() in _onFieldChanged with silent validator calls; removed errant onChanged callback.',
             '✅ Fixed'),
            ('Avatar path invalid after reinstall',
             'S2', 'Med',
             'Sprint 5 will migrate to content URI or base64. Within Sprint 2 path is used only within the active session.',
             '🔜 Sprint 5'),
            ('SHA-256 without salt (rainbow table)',
             'S2', 'Med',
             'Offline app — no network attack surface in Sprint 2. Salt + bcrypt recommended for Sprint 8 security hardening.',
             '🔜 Sprint 8'),
            ('image_picker permission denied',
             'S2', 'Low',
             'AndroidManifest already declares CAMERA. Graceful catch shown to user via SnackBar.',
             '✅ Mitigated'),
        ],
        col_widths=[1.6, 0.5, 0.5, 3.0, 0.9],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. SIGN-OFF
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '7. Sign-Off', 1)

    make_table(doc,
        headers=['Sprint', 'Role', 'Name', 'Status', 'Date'],
        rows=[
            ('Sprint 1', 'Developer', 'Mark Leannie Gacutno', '✅ Submitted', 'March 5, 2026'),
            ('Sprint 1', 'Reviewer',  'Tom (Team Lead)',       '☐ Pending',   '____________'),
            ('Sprint 1', 'QA',        'Rex (QA Lead)',         '☐ Pending',   '____________'),
            ('Sprint 2', 'Developer', 'Mark Leannie Gacutno', '✅ Submitted', 'March 5, 2026'),
            ('Sprint 2', 'Reviewer',  'Tom (Team Lead)',       '☐ Pending',   '____________'),
            ('Sprint 2', 'QA',        'Rex (QA Lead)',         '☐ Pending',   '____________'),
        ],
        col_widths=[0.8, 0.9, 1.8, 1.2, 1.8],
    )
    doc.add_paragraph()
    add_para(doc,
             'Sprint 3 (Portfolio & Projects CRUD) begins after both Sprint 1 and Sprint 2 are signed off.',
             italic=True, color=MID_GREY)

    # Footer
    doc.add_paragraph()
    hr_p = doc.add_paragraph()
    hr_p.paragraph_format.space_before = Pt(6)
    hr_r = hr_p.add_run('─' * 100)
    hr_r.font.size = Pt(8); hr_r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    foot_p = doc.add_paragraph()
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot_p.add_run(
        'PortFolioPH  |  Sprint 1 & Sprint 2 Combined Implementation Report  '
        '|  Mark Leannie Gacutno  |  March 5, 2026  '
        '|  github.com/auzcee/PortFolioPHH'
    )
    fr.font.size = Pt(8); fr.italic = True; fr.font.color.rgb = MID_GREY

    return doc


# ── Entry point ────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    out_dir  = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, 'PortFolioPH_Sprint1_Sprint2_Combined_Report.docx')

    document = build_document()
    document.save(out_path)

    abs_path = os.path.abspath(out_path)
    print(f'✅  Document saved : {abs_path}')
    print(f'   Sprints covered : Sprint 1 + Sprint 2')
    print(f'   Tasks detailed  : 8 (S1) + 11 (S2) + 2 bug fixes = 21 task cards')
    print(f'   Screenshot slots: ~35 Jira-ready placeholder sections')
