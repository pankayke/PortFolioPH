from docx import Document
from docx.shared import Pt


COLUMNS = [
    "Test ID",
    "Overview",
    "Precondition",
    "Test Data",
    "Test Steps",
    "Expected Result",
    "Actual Result",
    "Remarks",
]


def add_module_page(document: Document, module_title: str, cases: list[dict]) -> None:
    document.add_paragraph("USER: PORTFOLIOPH TEST CASE")
    document.add_paragraph(module_title)

    table = document.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, col in enumerate(COLUMNS):
        hdr[i].text = col

    for case in cases:
        row = table.add_row().cells
        row[0].text = case["id"]
        row[1].text = case["overview"]
        row[2].text = case["precondition"]
        row[3].text = case["test_data"]
        row[4].text = case["steps"]
        row[5].text = case["expected"]
        row[6].text = "To be executed"
        row[7].text = "PENDING"


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    modules = [
        (
            "SIGNUP MODULE",
            [
                {
                    "id": "PFVT-001",
                    "overview": "Open signup page and verify required fields render.",
                    "precondition": "User is on login page.",
                    "test_data": "None",
                    "steps": "Click Signup link.",
                    "expected": "Signup page shows Full Name, Email, Role, Password, Confirm Password, and Signup button.",
                },
                {
                    "id": "PFVT-002",
                    "overview": "Register with valid job seeker data.",
                    "precondition": "Signup page is open.",
                    "test_data": "name=Maria Cruz, email=maria.cruz@gmail.com, role=job_seeker, password=StrongPass123!",
                    "steps": "Fill all fields and submit.",
                    "expected": "Account is created and user is redirected to login/dashboard based on flow.",
                },
                {
                    "id": "PFVT-003",
                    "overview": "Register with duplicate email.",
                    "precondition": "Existing account uses target email.",
                    "test_data": "email=existing.user@gmail.com",
                    "steps": "Submit signup form using duplicate email.",
                    "expected": "Validation error is shown and account is not created.",
                },
                {
                    "id": "PFVT-004",
                    "overview": "Register with invalid email format.",
                    "precondition": "Signup page is open.",
                    "test_data": "email=invalid-email",
                    "steps": "Submit signup form.",
                    "expected": "Validation blocks request with email format error.",
                },
                {
                    "id": "PFVT-005",
                    "overview": "Register with weak password.",
                    "precondition": "Signup page is open.",
                    "test_data": "password=12345",
                    "steps": "Submit signup form.",
                    "expected": "Password policy error is shown and save is blocked.",
                },
                {
                    "id": "PFVT-006",
                    "overview": "Register with mismatched confirm password.",
                    "precondition": "Signup page is open.",
                    "test_data": "password=StrongPass123!, confirm=StrongPass124!",
                    "steps": "Submit signup form.",
                    "expected": "Mismatch error appears and account creation is blocked.",
                },
            ],
        ),
        (
            "LOGIN MODULE",
            [
                {
                    "id": "PFVT-007",
                    "overview": "Login page renders required controls.",
                    "precondition": "Application is loaded.",
                    "test_data": "None",
                    "steps": "Open Login page.",
                    "expected": "Email, Password, Login button, and Forgot Password option are visible.",
                },
                {
                    "id": "PFVT-008",
                    "overview": "Login with valid credentials.",
                    "precondition": "Valid account exists.",
                    "test_data": "email=valid.user@gmail.com, password=StrongPass123!",
                    "steps": "Enter credentials and click Login.",
                    "expected": "User is authenticated and routed to role-specific dashboard.",
                },
                {
                    "id": "PFVT-009",
                    "overview": "Login with incorrect password.",
                    "precondition": "Valid email exists.",
                    "test_data": "email=valid.user@gmail.com, password=WrongPass999",
                    "steps": "Submit login form.",
                    "expected": "Unauthorized error shown and session not created.",
                },
                {
                    "id": "PFVT-010",
                    "overview": "Login with unregistered email.",
                    "precondition": "No account for input email.",
                    "test_data": "email=missing.user@gmail.com",
                    "steps": "Submit login form.",
                    "expected": "Authentication fails with generic invalid credentials message.",
                },
                {
                    "id": "PFVT-011",
                    "overview": "Login with empty required fields.",
                    "precondition": "Login page is open.",
                    "test_data": "email=empty, password=empty",
                    "steps": "Click Login without filling fields.",
                    "expected": "Field-level validation errors are displayed.",
                },
                {
                    "id": "PFVT-012",
                    "overview": "Logout invalidates session/token.",
                    "precondition": "User is logged in.",
                    "test_data": "None",
                    "steps": "Click Logout from profile/menu.",
                    "expected": "User is redirected to login and protected routes become inaccessible.",
                },
            ],
        ),
        (
            "DASHBOARD MODULE",
            [
                {
                    "id": "PFVT-013",
                    "overview": "Load dashboard widgets by role.",
                    "precondition": "User is logged in.",
                    "test_data": "roles=admin,recruiter,job_seeker",
                    "steps": "Open dashboard for each role.",
                    "expected": "Role-specific cards and panels render correctly.",
                },
                {
                    "id": "PFVT-014",
                    "overview": "Search jobs from dashboard.",
                    "precondition": "Approved jobs exist.",
                    "test_data": "search=flutter",
                    "steps": "Enter term in search and submit.",
                    "expected": "List filters to matching jobs.",
                },
                {
                    "id": "PFVT-015",
                    "overview": "Search with special characters.",
                    "precondition": "Dashboard visible.",
                    "test_data": "search=flutter@#",
                    "steps": "Submit query.",
                    "expected": "No crash; sanitized search runs and returns safe result set.",
                },
                {
                    "id": "PFVT-016",
                    "overview": "Navigation links from dashboard shell.",
                    "precondition": "User is logged in.",
                    "test_data": "None",
                    "steps": "Open menu and click each destination.",
                    "expected": "Each route resolves and view loads with no authorization leak.",
                },
                {
                    "id": "PFVT-017",
                    "overview": "Pagination controls on dashboard lists.",
                    "precondition": "More than one page of data exists.",
                    "test_data": "per_page defaults",
                    "steps": "Navigate next/previous pages.",
                    "expected": "Records update correctly and page state remains consistent.",
                },
                {
                    "id": "PFVT-018",
                    "overview": "Realtime counters consistency.",
                    "precondition": "Known fixture data in DB.",
                    "test_data": "users/jobs/applications sample set",
                    "steps": "Compare visible counts against DB counts.",
                    "expected": "Dashboard metrics equal backend aggregates.",
                },
            ],
        ),
        (
            "PROFILE MODULE",
            [
                {
                    "id": "PFVT-019",
                    "overview": "Open profile page and verify editable fields.",
                    "precondition": "User is logged in.",
                    "test_data": "None",
                    "steps": "Navigate to profile.",
                    "expected": "Name, email, phone, bio, location, website and save controls are visible.",
                },
                {
                    "id": "PFVT-020",
                    "overview": "Update profile with valid text fields.",
                    "precondition": "Profile page open.",
                    "test_data": "name=Updated User, bio=Senior developer",
                    "steps": "Edit fields and save.",
                    "expected": "Changes persist and success feedback appears.",
                },
                {
                    "id": "PFVT-021",
                    "overview": "Upload valid avatar image.",
                    "precondition": "Profile page open.",
                    "test_data": "avatar=jpg/png under size limit",
                    "steps": "Select image and save profile.",
                    "expected": "Avatar uploads, path stored, and image renders in UI.",
                },
                {
                    "id": "PFVT-022",
                    "overview": "Upload unsupported avatar format.",
                    "precondition": "Profile page open.",
                    "test_data": "avatar=exe file",
                    "steps": "Attempt upload and save.",
                    "expected": "Validation rejects file type with safe error message.",
                },
                {
                    "id": "PFVT-023",
                    "overview": "Change email to existing account email.",
                    "precondition": "Another user already owns target email.",
                    "test_data": "email=duplicate@domain.com",
                    "steps": "Update email and submit.",
                    "expected": "Unique constraint validation error is shown.",
                },
                {
                    "id": "PFVT-024",
                    "overview": "Reset password from profile/security screen.",
                    "precondition": "Authenticated user session.",
                    "test_data": "new_password=StrongPass456!",
                    "steps": "Submit password reset form and re-login.",
                    "expected": "Password updates successfully and old password is invalid.",
                },
            ],
        ),
        (
            "RECRUITER JOB MANAGEMENT MODULE",
            [
                {
                    "id": "PFVT-025",
                    "overview": "Create job posting with valid payload.",
                    "precondition": "Logged in as recruiter.",
                    "test_data": "title, description, location, salary range, job_type, deadline",
                    "steps": "Open create job form, fill data, submit.",
                    "expected": "Job is saved and appears in recruiter job list.",
                },
                {
                    "id": "PFVT-026",
                    "overview": "Create job with missing required fields.",
                    "precondition": "Logged in as recruiter.",
                    "test_data": "title missing",
                    "steps": "Submit incomplete form.",
                    "expected": "422 validation errors are shown.",
                },
                {
                    "id": "PFVT-027",
                    "overview": "Job seeker attempts to create job.",
                    "precondition": "Logged in as job seeker.",
                    "test_data": "valid job payload",
                    "steps": "POST job create endpoint/form.",
                    "expected": "Access denied (403/redirect) and no job is created.",
                },
                {
                    "id": "PFVT-028",
                    "overview": "Recruiter updates own job.",
                    "precondition": "Recruiter owns existing job.",
                    "test_data": "title/status updates",
                    "steps": "Edit job and submit update.",
                    "expected": "Job updates and audit fields/timestamps refresh.",
                },
                {
                    "id": "PFVT-029",
                    "overview": "Recruiter attempts to update another recruiter's job.",
                    "precondition": "Recruiter logged in but does not own target job.",
                    "test_data": "job_id of another recruiter",
                    "steps": "Try edit/update request.",
                    "expected": "Authorization failure and no changes in DB.",
                },
                {
                    "id": "PFVT-030",
                    "overview": "Delete own job and validate cascade behavior.",
                    "precondition": "Job has linked applications.",
                    "test_data": "job with 1+ applications",
                    "steps": "Delete job from recruiter view.",
                    "expected": "Job is removed and dependent applications are removed/handled per policy.",
                },
            ],
        ),
        (
            "JOB APPLICATION MODULE",
            [
                {
                    "id": "PFVT-031",
                    "overview": "Job seeker submits application successfully.",
                    "precondition": "Approved job exists.",
                    "test_data": "job_id, cover_letter",
                    "steps": "Apply to job.",
                    "expected": "Application status is pending and visible in seeker history.",
                },
                {
                    "id": "PFVT-032",
                    "overview": "Prevent duplicate application for same job.",
                    "precondition": "Seeker already applied to target job.",
                    "test_data": "same job_id",
                    "steps": "Submit second application.",
                    "expected": "Duplicate attempt is blocked with proper error.",
                },
                {
                    "id": "PFVT-033",
                    "overview": "Apply to nonexistent job.",
                    "precondition": "Authenticated job seeker.",
                    "test_data": "job_id=999999",
                    "steps": "Submit application request.",
                    "expected": "Validation/not-found error returned and no record created.",
                },
                {
                    "id": "PFVT-034",
                    "overview": "Recruiter updates application status.",
                    "precondition": "Application belongs to recruiter-owned job.",
                    "test_data": "status=shortlisted",
                    "steps": "Open application and update status.",
                    "expected": "Status is updated and visible to applicant.",
                },
                {
                    "id": "PFVT-035",
                    "overview": "Job seeker cannot update application status.",
                    "precondition": "Authenticated job seeker session.",
                    "test_data": "status=accepted",
                    "steps": "Call status update action.",
                    "expected": "Forbidden response and status unchanged.",
                },
                {
                    "id": "PFVT-036",
                    "overview": "Application list filtering by status.",
                    "precondition": "User has applications in multiple statuses.",
                    "test_data": "status=pending",
                    "steps": "Load applications with status filter.",
                    "expected": "Only matching status rows are returned.",
                },
            ],
        ),
        (
            "ADMIN USERS MANAGEMENT MODULE",
            [
                {
                    "id": "PFVT-037",
                    "overview": "Admin users page loads with pagination and precomputed summary metrics.",
                    "precondition": "Logged in as active admin.",
                    "test_data": "20+ users in DB",
                    "steps": "Open admin users page.",
                    "expected": "Users table renders with pagination controls and sidebar metrics without recalculating counts in the view.",
                },
                {
                    "id": "PFVT-038",
                    "overview": "Users search supports multi-term matching.",
                    "precondition": "Admin users page loaded.",
                    "test_data": "search=alice dev",
                    "steps": "Apply search query.",
                    "expected": "Results match name/email/username across terms.",
                },
                {
                    "id": "PFVT-039",
                    "overview": "Role filter accepts job seeker aliases.",
                    "precondition": "Admin users page loaded.",
                    "test_data": "role=job seeker",
                    "steps": "Apply role filter.",
                    "expected": "Only job seeker users are listed.",
                },
                {
                    "id": "PFVT-040",
                    "overview": "Suspended status filter returns inactive users only.",
                    "precondition": "Users include active and inactive accounts.",
                    "test_data": "status=suspended",
                    "steps": "Apply status filter.",
                    "expected": "List contains only inactive/null-active users.",
                },
                {
                    "id": "PFVT-041",
                    "overview": "Sort by active desc prioritizes active users.",
                    "precondition": "Mixed active states in users table.",
                    "test_data": "sort_by=active, sort_dir=desc",
                    "steps": "Apply sorting.",
                    "expected": "Active users appear before inactive users.",
                },
                {
                    "id": "PFVT-042",
                    "overview": "Invalid sort key falls back to created_at desc.",
                    "precondition": "Known user creation timestamps.",
                    "test_data": "sort_by=not_a_column",
                    "steps": "Apply invalid sort key.",
                    "expected": "Newest users appear first with no SQL error.",
                },
            ],
        ),
        (
            "ADMIN JOBS MODERATION MODULE",
            [
                {
                    "id": "PFVT-043",
                    "overview": "Admin jobs page loads with recruiter relation and controller-provided summary metrics.",
                    "precondition": "Admin authenticated.",
                    "test_data": "jobs in approved/pending/closed/draft",
                    "steps": "Open admin jobs page.",
                    "expected": "Job rows and moderation counters render correctly without view-level collection scans.",
                },
                {
                    "id": "PFVT-044",
                    "overview": "Approve pending job.",
                    "precondition": "Pending job exists.",
                    "test_data": "job_id pending",
                    "steps": "Click approve action.",
                    "expected": "Job status changes to approved and success message appears.",
                },
                {
                    "id": "PFVT-045",
                    "overview": "Suspend approved job.",
                    "precondition": "Approved job exists.",
                    "test_data": "job_id approved",
                    "steps": "Click suspend/close action.",
                    "expected": "Job status changes to closed.",
                },
                {
                    "id": "PFVT-046",
                    "overview": "Delete job with applications uses safe cascade path.",
                    "precondition": "Target job has linked applications.",
                    "test_data": "job_id with dependents",
                    "steps": "Delete job from admin panel.",
                    "expected": "Job and linked applications are removed atomically.",
                },
                {
                    "id": "PFVT-047",
                    "overview": "Non-admin user denied admin jobs route.",
                    "precondition": "Logged in as recruiter/job seeker.",
                    "test_data": "None",
                    "steps": "Navigate to /admin/jobs.",
                    "expected": "Access denied and redirected away from admin route.",
                },
                {
                    "id": "PFVT-048",
                    "overview": "Admin job detail shows applications and recruiter info.",
                    "precondition": "Job exists with recruiter and applications.",
                    "test_data": "job_id",
                    "steps": "Open admin job detail page.",
                    "expected": "Recruiter metadata and applications list are visible.",
                },
            ],
        ),
        (
            "ADMIN APPLICATIONS ANALYTICS MODULE",
            [
                {
                    "id": "PFVT-049",
                    "overview": "Admin applications index loads with aggregate status stats.",
                    "precondition": "Applications exist in multiple statuses.",
                    "test_data": "pending/reviewed/shortlisted/accepted/rejected data",
                    "steps": "Open admin applications page.",
                    "expected": "Status cards match database counts.",
                },
                {
                    "id": "PFVT-050",
                    "overview": "Pagination works on admin applications table.",
                    "precondition": "More than 20 applications exist.",
                    "test_data": "default per-page",
                    "steps": "Navigate to page 2 and back.",
                    "expected": "Rows change by page and no duplicates across pages.",
                },
                {
                    "id": "PFVT-051",
                    "overview": "Each row resolves job and user relationship data.",
                    "precondition": "Applications have valid foreign keys.",
                    "test_data": "None",
                    "steps": "Inspect row details.",
                    "expected": "Job title and applicant information are displayed.",
                },
                {
                    "id": "PFVT-052",
                    "overview": "Admin sees pending queue count in settings metrics.",
                    "precondition": "Pending applications exist.",
                    "test_data": "pending records",
                    "steps": "Open settings page after creating pending applications.",
                    "expected": "Queue backlog/active sessions reflect pending count logic.",
                },
                {
                    "id": "PFVT-053",
                    "overview": "Analytics remains stable with zero application records.",
                    "precondition": "Applications table empty.",
                    "test_data": "None",
                    "steps": "Open applications index.",
                    "expected": "All stats display zero and page loads without errors.",
                },
                {
                    "id": "PFVT-054",
                    "overview": "Admin analytics route protected by admin middleware.",
                    "precondition": "Authenticated non-admin session.",
                    "test_data": "None",
                    "steps": "Attempt to open /admin/applications.",
                    "expected": "Request is denied and redirected with access warning.",
                },
            ],
        ),
        (
            "SECURITY, SETTINGS, AND SESSION MODULE",
            [
                {
                    "id": "PFVT-055",
                    "overview": "Admin settings save valid payload to session.",
                    "precondition": "Admin is authenticated.",
                    "test_data": "maintenance_mode=1, digest_frequency=weekly, dashboard_density=compact, session_timeout=45",
                    "steps": "Submit settings form.",
                    "expected": "Settings persist in session and success flash is shown.",
                },
                {
                    "id": "PFVT-056",
                    "overview": "Admin settings reject invalid payload.",
                    "precondition": "Admin is authenticated.",
                    "test_data": "digest_frequency=invalid, dashboard_density=ultra, session_timeout=1",
                    "steps": "Submit invalid settings form.",
                    "expected": "Validation errors appear and settings are not persisted.",
                },
                {
                    "id": "PFVT-057",
                    "overview": "Settings page uses default fallback values without session state.",
                    "precondition": "No admin_settings in session.",
                    "test_data": "None",
                    "steps": "Open settings page.",
                    "expected": "Defaults show: maintenance=false, new_user_alerts=true, moderation_alerts=true, digest=daily, density=high, timeout=30.",
                },
                {
                    "id": "PFVT-058",
                    "overview": "Suspended admin account is logged out by middleware.",
                    "precondition": "Admin account active=false.",
                    "test_data": "None",
                    "steps": "Login and access admin route.",
                    "expected": "Session is terminated and user is redirected to login with suspension notice.",
                },
                {
                    "id": "PFVT-059",
                    "overview": "Public jobs endpoint clamps oversized per_page values.",
                    "precondition": "Approved jobs exist and public API is reachable.",
                    "test_data": "GET /api/jobs?per_page=500",
                    "steps": "Call jobs index with an oversized per_page value.",
                    "expected": "Response returns at most 100 jobs and remains paginated.",
                },
                {
                    "id": "PFVT-060",
                    "overview": "Recruiter jobs endpoint clamps per_page and remains routable.",
                    "precondition": "Recruiter is authenticated and owns jobs.",
                    "test_data": "GET /api/jobs/mine?per_page=500",
                    "steps": "Call recruiter jobs list with an oversized per_page value.",
                    "expected": "Response returns at most 100 recruiter jobs and resolves the mine route correctly.",
                },
                {
                    "id": "PFVT-061",
                    "overview": "Public job detail page loads recruiter data and paginated applications instead of the full relation.",
                    "precondition": "Recruiter is authenticated and a job has more than ten applications.",
                    "test_data": "GET /jobs/{id} with 12 applications",
                    "steps": "Open the job detail page as the owning recruiter.",
                    "expected": "Applications render from the paginated collection, the counter shows the total, and the view avoids loading the entire relation.",
                },
            ],
        ),
    ]

    for idx, (module_title, cases) in enumerate(modules):
        add_module_page(doc, module_title, cases)
        if idx < len(modules) - 1:
            doc.add_page_break()

    output_path = "docs/PortFolioPH_10_Page_Test_Cases.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
